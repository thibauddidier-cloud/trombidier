#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de test pour la génération de trombinoscope (sans GUI)
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def sort_class_name(class_name):
    """Tri intelligent des noms de classes"""
    if class_name.startswith('2DE'):
        number = int(re.search(r'\d+', class_name).group())
        return (1, number, class_name)
    elif class_name.startswith('PG'):
        number = int(re.search(r'\d+', class_name).group())
        return (2, number, class_name)
    elif class_name.startswith('PSTMG'):
        number = int(re.search(r'\d+', class_name).group())
        return (3, number, class_name)
    elif class_name.startswith('TG') and not class_name.startswith('TGF'):
        number = int(re.search(r'\d+', class_name).group())
        return (4, number, class_name)
    elif class_name.startswith('TM') or class_name.startswith('TGF') or class_name.startswith('TRHC'):
        match = re.search(r'\d+', class_name)
        number = int(match.group()) if match else 0
        return (5, number, class_name)
    elif class_name.startswith('BTS'):
        bts_match = re.search(r'BTS\s*([A-Z]+)(\d+)', class_name)
        if bts_match:
            bts_type = bts_match.group(1)
            number = int(bts_match.group(2))
            return (6, ord(bts_type[0]), number, class_name)
        return (6, 0, 0, class_name)
    else:
        return (99, 0, class_name)


def analyze_classes(folder_path):
    """Analyse les classes dans le dossier"""
    classes_data = {}
    
    subdirs = [d for d in os.listdir(folder_path) 
              if os.path.isdir(os.path.join(folder_path, d))]
    
    subdirs.sort(key=sort_class_name)
    
    for class_name in subdirs:
        class_path = os.path.join(folder_path, class_name)
        students = []
        
        for file in os.listdir(class_path):
            if file.lower().endswith(('.jpg', '.jpeg', '.png')):
                name = os.path.splitext(file)[0]
                students.append({
                    'name': name,
                    'file_path': os.path.join(class_path, file)
                })
        
        students.sort(key=lambda x: x['name'])
        
        if students:
            classes_data[class_name] = students
    
    return classes_data


def create_word_document(classes_data, output_file, school_name="Lycée Toulouse Lautrec", school_year="2024-2025"):
    """Création du document Word"""
    doc = Document()
    
    # Configuration de la page en paysage
    for section in doc.sections:
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Page de couverture
    add_cover_page(doc, classes_data, school_name, school_year)
    
    # Pages des classes
    for idx, (class_name, students) in enumerate(classes_data.items()):
        if idx > 0:
            doc.add_page_break()
        add_class_page(doc, class_name, students, school_name, school_year)
    
    doc.save(output_file)
    print(f"✅ Document généré : {output_file}")


def add_cover_page(doc, classes_data, school_name, school_year):
    """Ajout de la page de couverture"""
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = header.add_run(school_name)
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)
    
    year_para = doc.add_paragraph()
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year_para.add_run(f"\nAnnée Scolaire {school_year}")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(5, 150, 105)
    
    doc.add_paragraph("\n" * 3)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TROMBINOSCOPE")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)
    
    doc.add_paragraph("\n" * 5)
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"{len(classes_data)} Classes")
    run.font.size = Pt(18)
    
    total_students = sum(len(students) for students in classes_data.values())
    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info2.add_run(f"{total_students} Élèves")
    run.font.size = Pt(18)


def add_class_page(doc, class_name, students, school_name, school_year):
    """Ajout d'une page de classe"""
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = header.add_run(f"{school_name}  •  ")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(30, 58, 138)
    
    run = header.add_run(f"{school_year}  •  ")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(5, 150, 105)
    
    run = header.add_run(class_name)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)
    
    doc.add_paragraph()
    
    num_students = len(students)
    
    if num_students <= 28:
        rows = 4
        cols = 7
    elif num_students <= 35:
        rows = 5
        cols = 7
    else:
        rows = 6
        cols = 6
    
    if rows == 6:
        photo_width = Cm(3.2)
    else:
        photo_width = Cm(3.5)
    
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for idx, student in enumerate(students):
        row = idx // cols
        col = idx % cols
        
        if row >= rows:
            break
        
        cell = table.rows[row].cells[col]
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.add_run()
            run.add_picture(student['file_path'], width=photo_width)
            
            name_parts = student['name'].split('-')
            if len(name_parts) == 2:
                prenom = name_parts[1].capitalize()
                nom = name_parts[0].upper()
                display_name = f"{prenom}\n{nom}"
            else:
                display_name = student['name']
            
            name_para = cell.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = name_para.add_run(display_name)
            run.font.size = Pt(8)
            run.font.bold = True
            
        except Exception as e:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(f"[Photo manquante]\n{student['name']}")
            run.font.size = Pt(8)


def main():
    """Test de l'application"""
    print("="*60)
    print("🧪 TEST DE L'APPLICATION TROMBINOSCOPE")
    print("="*60)
    
    # Test avec les données d'exemple
    sample_folder = "/app/sample_data/CLASSE_JPG"
    
    if not os.path.exists(sample_folder):
        print("❌ Dossier d'exemple non trouvé")
        return
    
    print(f"\n📁 Analyse du dossier : {sample_folder}")
    classes_data = analyze_classes(sample_folder)
    
    print(f"\n✅ {len(classes_data)} classe(s) trouvée(s):")
    
    total_students = 0
    for class_name, students in classes_data.items():
        print(f"   • {class_name}: {len(students)} élève(s)")
        for student in students:
            parts = student['name'].split('-')
            if len(parts) == 2:
                prenom = parts[1].capitalize()
                nom = parts[0].upper()
                print(f"      - {prenom} {nom}")
            else:
                print(f"      - {student['name']}")
        total_students += len(students)
    
    print(f"\n📊 Total : {total_students} élève(s)")
    
    # Génération du document
    output_file = "/app/Trombinoscope_Test.docx"
    print(f"\n📄 Génération du document Word...")
    create_word_document(classes_data, output_file)
    
    print(f"\n✅ Test terminé avec succès!")
    print(f"📄 Document généré : {output_file}")
    print("\n" + "="*60)


if __name__ == "__main__":
    main()
