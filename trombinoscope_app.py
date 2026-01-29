#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Application de Génération de Trombinoscope
Lycée Toulouse Lautrec
Version 2.0 - Corrigée par Thibaud DIDIER
Version 2.5 - Casino Animé Edition
"""

import os
import re
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from PIL import Image, ImageTk
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil
from datetime import datetime
import random
import json
import threading
import math

# Tentative d'importer pygame pour les sons (optionnel)
try:
    import pygame
    pygame.mixer.init(frequency=22050, size=-16, channels=2, buffer=512)
    PYGAME_AVAILABLE = True
except ImportError:
    PYGAME_AVAILABLE = False


class TrombinoscopeApp:
    """Application principale de génération de trombinoscope"""
    
    # Chemin du fichier de sauvegarde pour les phrases achetées
    SAVE_FILE = os.path.join(os.path.dirname(__file__), "casino_save.json")
    
    def __init__(self, root):
        self.root = root
        self.root.title("Générateur de Trombinoscope - Lycée Toulouse Lautrec")
        self.root.geometry("900x750")
        
        # Variables
        self.classe_jpg_path = tk.StringVar()
        self.output_format = tk.StringVar(value="word")
        self.school_name = tk.StringVar(value="Lycée Toulouse Lautrec")
        self.school_year = tk.StringVar(value="2024-2025")
        self.classes_data = {}
        
        # Charger les phrases achetées depuis le fichier de sauvegarde
        self.purchased_phrases = self.load_purchased_phrases()
        
        # Charger les phrases personnalisées ajoutées par l'utilisateur
        self.load_custom_phrases()
        
        # Couleurs institutionnelles
        self.color_blue = "#000000"  # Noir (bandeau)
        self.color_green = "#059669"  # Vert
        self.color_light_blue = "#3b82f6"  # Bleu clair
        self.color_bg = "#f0f9ff"  # Fond bleu très clair
        
        # Messages Easter Egg
        self.easter_messages = [
            "La bouffe du self ou plutot de la cuisine centrale est dégueulasse. Et le prix va augmenter. 2'80e a 3'25e en 1 an et imaginez les externes qui payent plus de 9e",
            "Le saviez-vous ? Sophian est l'AED ayant passé le plus de temps dans mon bureau, estimant que c'est sa planque.",
            "Spéciale dédicace à Karen qui anime la vie lycéenne !",
"La légende raconte qu'un certain Quentin Monnier serait directeur de colo ( et pas que ! )",
"Christophe Pagis est l'enseignant avec les plus belles cravattes du lycée.",
"Jarod peut remonter le temps grace à sa montre pendentif, seulement pendant son sommeil.",
"Un certain S est en concurence avec C pour la séduction des enseignants.",
"Larbi, ancien agent à Lautrec, a voulu se battre avec l'AED TICE et s'est fait virer depuis",
"Les cookies de la cafet sont délicieux. Ne prenez pas le tacos crousty il est surcoté ",
"Chose la plus insolite retrouvée dans un casier de prof : Un trognon de pomme et un dentier.",
"Non Sarah, ce n'est pas à moi de changer la cartouche de ton imprimante.",
"Les élèves de Mr Tempier semblent avoir le droit de faire un baseball en plein cours.",
"Regardez le top 10 des élèves sanctionnés sur pronote, c'est surprenant. ",
"En 2025, 2 AED ont vécu des choses pationnantes dans une salle de cours. Plusieurs fois. En salle 154 !",
"Mr Maiurano, ancien CPE, avait comme mot de passe pomme de pain ",
"Cathy Condy a été ma maman spirituelle. Pauline Ndoume a pris la relève, et Monique Fernandez est ma maîtresse",
"Clement est un futur pompier, qui n'éteint pas que le feu des batiments.",
"Non ECE, je ne te donnerai pas les identifiants admin pour que tu installes ton putain de logiciel Paint.net ",
"Karen a entendu ' Il en a une grosse ' à la place de 'il a des gosses'.",
"Les talky walky des AED captaient la fréquence du chantier d'a coté et c'etait très drole .",
"L'un des anciens AED, Samir, qui est resté que quelques semaines, aurait fumé de la marie-jeanne avec des élèves. Il été moniteur de karaté pour les AED aussi",
"Marc Andral est un énorme enfoiré, qui ne m'a donné aucune aide et adorait me faire passer pour un bouffon.",
"Règle 1, si Marie Hélène appelle de la part de Sarah, Sarah veut te douiller.",
"Règle 2, si Sarah appelle pour te demander un service, Sarah veut te douiller.",
"Jarod, ta bien-pensance est remarquable, ta non-tolérance l'est moins, mais je t'aime. Et non je ne suis pas d'Extreme droite! Vive l'UPR",
"Maryse fait la gueule toute la matinée ( sauf si Clément est présent ).",
"Maryse et Clément ne se parlent plus depuis des semaines, je suis sûr que l'un des deux à des sentiments envers l'autre.",
"Une AED a déjà fait un resto basket au 100 couvert pour la saint valentin !",
"Une pièce souterraine avec des poupées et des rituels vaudous a existé au lycée",
"Le tacos aux légumes de la cafet est sous-coté mais il est très bon.",
"Julien Pharos fait pleuvoir lorsqu'il chante au Boulis.",
"J'ai fait un énorme travail personnel pour apprécier ECE, ne supportant pas le ton de sa voix et ses demandes incessantes. Au final c'est une très bonne personne. Travaillez sur vous même c'est important.",
"Quoi que je dise sur Sophian pour le taquiner, c'est celui en qui j'ai trouvé la plus grande sagesse, neutralité et bienveillance. C'est juste un mec Chill.",
"Chaque année, des dizaines de rats crevés sont découverts dans les faux-plafonds, spécifiquement au dessus des vies scolaires.",
"En Octobre 2025, quelqu'un a chié au niveau de la passerelle #cacagate",
"Des lacrymos ont été entierement vidées dans les WC en 2025",
"50% des demandes des profs se terminent par 'fais le quand tu as le temps, ce n'est pas urgent' mais sont relancées dans la journée.",
"Cheveux de feu s'est battue avec sa fille dans la cour du lycée en faisant des roulé-boulé dans l'herbe, vue par des AED.",
"Des AED qui chopent des profs ? Oui ça s'est déjà fait. Des AED qui chopent des AED ? Oui, ça s'est déjà fait. 'Le Chene est dans la Scierie'. En salle 154 ! Le pire c'est que même les profs de ces salles ont été 'fierement' informé par l'une des AED!",
"Un élève a donné comme devoir PETE MOI LA CHATTE , en se connectant au Pronote d'une prof.",
"La prof qui s'est fait 'pirater' son Pronote, a finalement admis qu'elle s'était absentée 15mn du cours.",
"Marc Andral, oui encore, a dit devant tout le monde que j'étais au niveau -1 en informatique, pour au final ne pas savoir résoudre le problème pour lequel je l'avais été appelé.",
"Julien Pharos est l'enseignant avec le plus de retard au lycée, environ 15 mn par cours.",
"Pauline , l'agent d'entretien a été ma seconde maman spirituelle, hein St Antoine de Padou !",
"Certains contacts du téléphone pro ont été renommé avant l'arrivée de Paul. Mais Paul va surrement renommer certains contacts aussi",
"Mme Novoa est objectivement l'une des personnes les plus humaines de l'établissement. ",
"Pour votre santé mentale, ne demandez pas d'explications sur un problème simple à Michel Chaboy",
"Julien pharos a reçu un tracte d'Eric Zemmour dans son casier par pure provocation.",
"Marc Andral a été la personne la moins pédagogique de l'établissement, comble d'un prof.",
"Edouard Romera & Mylene Fournier sont les profs en couple les plus cools du lycée !",
"Le bourrage papier de l'imprimante de Davina était en fait un sachet de beuh coincé. Je savais pas d'où venait cette odeur de merde jusqu'à ce que je comprenne que ça vienne de cet emballage de 'gateau' mdrrr",
"Brigitte de la Cafet a écrit un livre. J'aurai pu lire son livre plus rapidement que tout le temps passé accumulé à la dépanner pour ses soucis persos #humour",
"Mr Audouard récupère des cartons du lycée pour son garage et son poulailler .",
"Sirine est l'AED a qui vous ne pouvez pas confier aucun secret. J'en ai pas fait les frais mais j'aurai pu lancer une rumeur et voir combien de temps elle met à faire le tour des AED avant de me revenir",
"Sophian investi dans des paris sportifs grâce à ChatGPT, mais reste en négatif.",
"Je me suis retenu de leak le 06 de Marc Andral sur Leboncoin."
        ]
        self.easter_message_index = 0  # Index pour messages séquentiels
        
        # Messages de chargement aléatoires
        self.loading_messages = [
            "Ajout d'accusation des mérovingiens",
            "Suppression des organes de Marc Andral",
            "Correction colorimétrique de Cheveux de feu",
            "Redimensionnement de MimiMathy.png",
            "Ajout de pépites de chocolat sur les photos de classe",
            "Conversion des photos au format .fdp",
            "Mise en surtension de Michel Chaboy",
            "Téléchargement du DLC « France du tiers monde »",
            "Le saviez-vous ? Les pieuvres ont trois cœurs et leur sang est bleu.",
            "Le saviez-vous ? Les vaches ont des meilleures amies et stressent quand elles sont séparées.",
            "Le saviez-vous ? Les escargots peuvent dormir jusqu'à trois ans.",
            "Le saviez-vous ? Les manchots demandent leur partenaire en mariage avec un caillou.",
            "Le saviez-vous ? Les requins ne peuvent pas tomber malades du cancer.",
            "Le saviez-vous ? Au Japon, il existe des îles entières peuplées de chats.",
            "Le saviez-vous ? La France possède le plus grand nombre de fuseaux horaires au monde.",
            "Le saviez-vous ? En Écosse, il existe 421 mots pour dire « neige ».",
            "Le saviez-vous ? En Australie, il y a plus de kangourous que d'humains.",
            "Le saviez-vous ? En Bolivie, il existe un marché de sorcellerie totalement légal.",
            "Le saviez-vous ? Au Danemark, on casse de la vaisselle devant la porte des amis pour leur porter chance.",
            "Le saviez-vous ? En Corée du Sud, l'âge commence traditionnellement à un an dès la naissance.",
            "Le saviez-vous ? En allemand, il existe un mot pour le plaisir ressenti face à l'échec d'autrui : Schadenfreude.",
            "Le saviez-vous ? Les grenouilles peuvent geler complètement en hiver… puis revenir à la vie au printemps.",
            "Le saviez-vous ? Les vaches peuvent monter les escaliers mais pas les descendre.",
            "Le saviez-vous ? Les dauphins dorment avec un seul hémisphère du cerveau à la fois.",
            "Le saviez-vous ? En Écosse, on peut frapper à n'importe quelle porte le Nouvel An pour porter chance.",
            "Le saviez-vous ? En France, il est illégal de nommer un cochon Napoléon.",
            "Le saviez-vous ? Les humains partagent environ 60 % de leur ADN avec les bananes.",
            "Le saviez-vous ? Certains champignons peuvent contrôler le cerveau des insectes.",
            "Le saviez-vous ? Les axolotls peuvent régénérer leur cerveau et leur cœur.",
            "Le saviez-vous ? Les crevettes-pistolets produisent un son plus fort qu'un coup de feu.",
            "Le saviez-vous ? Les requins peuvent vivre plus de 400 ans (requin du Groenland).",
            "Le saviez-vous ? La Tour Eiffel grandit d'environ 15 cm en été.",
            "Le saviez-vous ? Ton nez peut reconnaître plus d'un trillion d'odeurs."
        ]
        
        self.setup_ui()
        self.create_menu()
    
    def create_modern_button(self, parent, text, command, bg_color, width=None):
        """Créer un bouton moderne avec coins arrondis et style amélioré"""
        # Frame pour simuler les coins arrondis avec un effet d'ombre
        button_container = tk.Frame(parent, bg=parent.cget('bg'))
        
        button = tk.Button(
            button_container,
            text=text,
            command=command,
            bg=bg_color,
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=12,
            activebackground=self.get_darker_color(bg_color),
            activeforeground="white",
            borderwidth=0,
            highlightthickness=0
        )
        
        if width:
            button.config(width=width)
        
        button.pack()
        
        # Effets hover
        def on_enter(e):
            button.config(bg=self.get_darker_color(bg_color))
        
        def on_leave(e):
            button.config(bg=bg_color)
        
        button.bind("<Enter>", on_enter)
        button.bind("<Leave>", on_leave)
        
        return button_container
    
    def get_darker_color(self, hex_color):
        """Obtenir une version plus foncée d'une couleur"""
        # Supprimer le #
        hex_color = hex_color.lstrip('#')
        # Convertir en RGB
        r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        # Assombrir de 20%
        r = max(0, int(r * 0.8))
        g = max(0, int(g * 0.8))
        b = max(0, int(b * 0.8))
        # Reconvertir en hex
        return f'#{r:02x}{g:02x}{b:02x}'
        
    def create_menu(self):
        """Création du menu avec l'onglet À propos, Aide et Extra"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # Menu Aide
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Aide", menu=help_menu)
        help_menu.add_command(label="Guide d'utilisation", command=self.show_help)
        
        # Menu À propos
        about_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="À propos", menu=about_menu)
        about_menu.add_command(label="À propos de l'application", command=self.show_about)
        
        # Menu Extra (Easter Eggs et fonctionnalités bonus)
        extra_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Extra", menu=extra_menu)
        extra_menu.add_command(label="🦆 Tape le Psykokwak", command=self.show_whack_a_psyduck_game)
        extra_menu.add_command(label="🎰 Le Casino de Lautrec", command=self.show_casino_game)
        extra_menu.add_separator()
        extra_menu.add_command(label="📝 Ajouter des phrases", command=self.show_add_phrases_window)
    
    def load_purchased_phrases(self):
        """Charger les phrases achetées et les crédits depuis le fichier de sauvegarde"""
        try:
            if os.path.exists(self.SAVE_FILE):
                with open(self.SAVE_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # Charger aussi les crédits sauvegardés
                    self.saved_credits = data.get('casino_credits', 100)
                    return set(data.get('purchased_phrases', []))
        except Exception as e:
            print(f"Erreur chargement sauvegarde: {e}")
        self.saved_credits = 100  # Valeur par défaut
        return set()
    
    def load_custom_phrases(self):
        """Charger les phrases personnalisées depuis le fichier JSON"""
        try:
            if os.path.exists(self.SAVE_FILE):
                with open(self.SAVE_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    custom_phrases = data.get('custom_phrases', [])
                    # Ajouter les phrases personnalisées si elles n'existent pas déjà
                    for phrase in custom_phrases:
                        if phrase not in self.easter_messages:
                            self.easter_messages.append(phrase)
                    if custom_phrases:
                        print(f"📜 {len(custom_phrases)} phrases personnalisées chargées")
        except Exception as e:
            print(f"Erreur chargement phrases personnalisées: {e}")
    
    def save_purchased_phrases(self):
        """Sauvegarder les phrases achetées et les crédits dans le fichier"""
        try:
            # Récupérer les crédits actuels du casino si disponible
            credits = getattr(self, 'casino_state', {}).get('credits', self.saved_credits)
            
            data = {
                'purchased_phrases': list(self.purchased_phrases),
                'casino_credits': credits
            }
            with open(self.SAVE_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"💾 Sauvegarde: {credits} crédits")
        except Exception as e:
            print(f"Erreur sauvegarde: {e}")
    
    def play_sound(self, sound_type):
        """Jouer un effet sonore rétro style Game Boy"""
        if not PYGAME_AVAILABLE:
            return
        
        def _play():
            try:
                # Créer des sons 8-bit style Game Boy
                sample_rate = 22050
                
                if sound_type == 'spin':
                    # Son de démarrage des rouleaux - bruit montant
                    duration = 0.15
                    freq_start, freq_end = 200, 600
                    samples = int(sample_rate * duration)
                    sound_array = []
                    for i in range(samples):
                        freq = freq_start + (freq_end - freq_start) * (i / samples)
                        val = int(127 * (1 if (i * freq / sample_rate) % 1 < 0.5 else -1))
                        sound_array.append(val)
                    
                elif sound_type == 'reel_stop':
                    # Son d'arrêt d'un rouleau - blip court
                    duration = 0.08
                    freq = 440
                    samples = int(sample_rate * duration)
                    sound_array = []
                    for i in range(samples):
                        decay = 1 - (i / samples)
                        val = int(100 * decay * (1 if (i * freq / sample_rate) % 1 < 0.5 else -1))
                        sound_array.append(val)
                    
                elif sound_type == 'win':
                    # Son de victoire - mélodie montante
                    duration = 0.4
                    samples = int(sample_rate * duration)
                    sound_array = []
                    notes = [523, 659, 784, 1047]  # Do Mi Sol Do (octave)
                    note_len = samples // len(notes)
                    for note_idx, freq in enumerate(notes):
                        for i in range(note_len):
                            val = int(100 * (1 if ((note_idx * note_len + i) * freq / sample_rate) % 1 < 0.5 else -1))
                            sound_array.append(val)
                    
                elif sound_type == 'jackpot':
                    # Son de jackpot - fanfare
                    duration = 0.6
                    samples = int(sample_rate * duration)
                    sound_array = []
                    notes = [523, 523, 523, 659, 784, 784, 659, 784, 1047]
                    note_len = samples // len(notes)
                    for note_idx, freq in enumerate(notes):
                        for i in range(note_len):
                            val = int(110 * (1 if ((note_idx * note_len + i) * freq / sample_rate) % 1 < 0.5 else -1))
                            sound_array.append(val)
                    
                elif sound_type == 'lose':
                    # Son de perte - ton descendant triste
                    duration = 0.3
                    freq_start, freq_end = 400, 150
                    samples = int(sample_rate * duration)
                    sound_array = []
                    for i in range(samples):
                        freq = freq_start + (freq_end - freq_start) * (i / samples)
                        decay = 1 - (i / samples) * 0.5
                        val = int(80 * decay * (1 if (i * freq / sample_rate) % 1 < 0.5 else -1))
                        sound_array.append(val)
                    
                elif sound_type == 'purchase':
                    # Son d'achat - cha-ching!
                    duration = 0.25
                    samples = int(sample_rate * duration)
                    sound_array = []
                    notes = [880, 1109, 1319]  # La Do# Mi
                    note_len = samples // len(notes)
                    for note_idx, freq in enumerate(notes):
                        for i in range(note_len):
                            val = int(90 * (1 if ((note_idx * note_len + i) * freq / sample_rate) % 1 < 0.5 else -1))
                            sound_array.append(val)
                    
                elif sound_type == 'error':
                    # Son d'erreur - buzzer
                    duration = 0.2
                    freq = 150
                    samples = int(sample_rate * duration)
                    sound_array = []
                    for i in range(samples):
                        val = int(70 * (1 if (i * freq / sample_rate) % 1 < 0.5 else -1))
                        sound_array.append(val)
                else:
                    return
                
                # Créer le son avec pygame
                import numpy as np
                sound_array = np.array(sound_array, dtype=np.int8)
                # Convertir en stéréo
                stereo = np.column_stack((sound_array, sound_array))
                sound = pygame.sndarray.make_sound(stereo)
                sound.play()
                
            except Exception as e:
                print(f"Erreur son: {e}")
        
        # Jouer le son dans un thread séparé pour ne pas bloquer l'UI
        threading.Thread(target=_play, daemon=True).start()
        
    def show_about(self):
        """Afficher la fenêtre À propos"""
        messagebox.showinfo(
            "À propos",
            "Ce logiciel a été réalisé par Thibaud DIDIER,\n"
            "pour aider les prochains AED TICE esclavagisés\n"
            "par ces tâches ingrates."
        )
    
    def show_help(self):
        """Afficher le guide d'aide depuis README_TROMBINOSCOPE.md"""
        readme_path = os.path.join(os.path.dirname(__file__), "README_TROMBINOSCOPE.md")
        
        # Créer une fenêtre popup
        popup = tk.Toplevel(self.root)
        popup.title("📚 Guide d'utilisation - Trombinoscope")
        popup.geometry("900x700")
        popup.configure(bg="white")
        
        # Centrer la fenêtre
        popup.transient(self.root)
        
        # Frame principale
        main_frame = tk.Frame(popup, bg="white", padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Titre
        title_label = tk.Label(
            main_frame,
            text="📚 Guide d'utilisation",
            font=("Arial", 16, "bold"),
            bg="white",
            fg=self.color_blue
        )
        title_label.pack(pady=(0, 15))
        
        # Zone de texte avec défilement pour afficher le README
        text_frame = tk.Frame(main_frame, bg="white")
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        help_text = scrolledtext.ScrolledText(
            text_frame,
            font=("Courier", 9),
            wrap=tk.WORD,
            bg="#f8fafc",
            relief=tk.FLAT,
            padx=10,
            pady=10
        )
        help_text.pack(fill=tk.BOTH, expand=True)
        
        # Charger et afficher le contenu du README
        try:
            if os.path.exists(readme_path):
                with open(readme_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    help_text.insert(tk.END, content)
            else:
                help_text.insert(tk.END, "⚠️ Fichier README_TROMBINOSCOPE.md introuvable.\n\n")
                help_text.insert(tk.END, "Le guide d'utilisation devrait se trouver dans le dossier de l'application.")
        except Exception as e:
            help_text.insert(tk.END, f"❌ Erreur lors du chargement du guide :\n{str(e)}")
        
        help_text.config(state=tk.DISABLED)
        
        # Bouton Fermer
        close_container = tk.Frame(main_frame, bg="white")
        close_container.pack(pady=(15, 0))
        
        close_btn = tk.Button(
            close_container,
            text="Fermer",
            command=popup.destroy,
            bg=self.color_green,
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=30,
            pady=10,
            activebackground="#047857",
            borderwidth=0,
            highlightthickness=0
        )
        close_btn.pack()
        
        def close_on_enter(e):
            close_btn.config(bg="#047857")
        def close_on_leave(e):
            close_btn.config(bg=self.color_green)
        close_btn.bind("<Enter>", close_on_enter)
        close_btn.bind("<Leave>", close_on_leave)
        
    def show_easter_egg(self):
        """Afficher un message Easter Egg séquentiel (appelé après victoire du mini-jeu)"""
        # Utiliser l'index actuel au lieu de random
        message = self.easter_messages[self.easter_message_index]
        
        # Incrémenter l'index pour le prochain clic (avec boucle)
        self.easter_message_index = (self.easter_message_index + 1) % len(self.easter_messages)
        
        return message  # Retourner le message pour l'utiliser dans le mini-jeu
    
    def show_whack_a_psyduck_game(self):
        """Afficher le mini-jeu Tape-Taupe avec Psykokwak"""
        # Créer une fenêtre popup pour le jeu
        game_popup = tk.Toplevel(self.root)
        game_popup.title("🎮 Assomme le Psykokwak !")
        game_popup.geometry("700x750")
        game_popup.configure(bg="#1a1a2e")
        game_popup.resizable(False, False)
        
        # Centrer la fenêtre
        game_popup.transient(self.root)
        game_popup.grab_set()
        
        # Variables du jeu stockées dans self pour accès global
        self.game_state = {
            'hits': 0,
            'target_hits': 3,
            'current_hole': None,
            'game_active': True,
            'psyduck_visible': False,
            'game_timer': None,
            'popup': game_popup,
            'hole_labels': [],
            'score_label': None,
            'result_label': None,
            'replay_btn': None,
            'close_btn': None,
            'psyduck_image': None
        }
        
        # Frame principale avec fond dégradé simulé
        main_frame = tk.Frame(game_popup, bg="#1a1a2e", padx=30, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Titre du jeu avec style
        title_frame = tk.Frame(main_frame, bg="#1a1a2e")
        title_frame.pack(pady=(0, 10))
        
        title_label = tk.Label(
            title_frame,
            text="🦆 ASSOMME LE PSYKOKWAK ! 🦆",
            font=("Arial", 20, "bold"),
            bg="#1a1a2e",
            fg="#ffd700"
        )
        title_label.pack()
        
        subtitle_label = tk.Label(
            title_frame,
            text="Clique sur Psykokwak 3 fois pour obtenir un secret !",
            font=("Arial", 11),
            bg="#1a1a2e",
            fg="#a0a0a0"
        )
        subtitle_label.pack(pady=(5, 0))
        
        # Compteur de score
        score_frame = tk.Frame(main_frame, bg="#16213e", padx=20, pady=10)
        score_frame.pack(pady=(10, 20))
        
        self.game_state['score_label'] = tk.Label(
            score_frame,
            text="🎯 Score: 0 / 3",
            font=("Arial", 16, "bold"),
            bg="#16213e",
            fg="#00ff88"
        )
        self.game_state['score_label'].pack()
        
        # Frame du jeu (grille 3x3)
        game_frame = tk.Frame(main_frame, bg="#1a1a2e")
        game_frame.pack(pady=10)
        
        # Charger l'image du psyduck pour le jeu
        try:
            psyduck_game_path = os.path.join(os.path.dirname(__file__), "assets", "psyduck.png")
            psyduck_game_img = Image.open(psyduck_game_path)
            psyduck_game_img = psyduck_game_img.resize((70, 70), Image.Resampling.LANCZOS)
            self.game_state['psyduck_image'] = ImageTk.PhotoImage(psyduck_game_img)
            
            # Charger l'image psyduck-ko pour l'animation de frappe
            psyduck_ko_path = os.path.join(os.path.dirname(__file__), "assets", "psyduck-ko.png")
            psyduck_ko_img = Image.open(psyduck_ko_path)
            psyduck_ko_img = psyduck_ko_img.resize((70, 70), Image.Resampling.LANCZOS)
            self.game_state['psyduck_ko_image'] = ImageTk.PhotoImage(psyduck_ko_img)
        except Exception as e:
            print(f"Erreur chargement psyduck pour le jeu: {e}")
        
        # Créer les 9 trous (3x3)
        for row in range(3):
            for col in range(3):
                # Container pour chaque trou avec effet d'ombre
                hole_container = tk.Frame(game_frame, bg="#1a1a2e", padx=8, pady=8)
                hole_container.grid(row=row, column=col)
                
                # Le "trou" avec effet 3D
                hole_frame = tk.Frame(
                    hole_container,
                    bg="#0d1117",
                    highlightbackground="#2d3748",
                    highlightthickness=3,
                    width=100,
                    height=100
                )
                hole_frame.pack_propagate(False)
                hole_frame.pack()
                
                # Label pour afficher le psyduck ou rien
                hole_label = tk.Label(
                    hole_frame,
                    bg="#0d1117",
                    width=100,
                    height=100,
                    cursor="crosshair"
                )
                hole_label.pack(expand=True, fill=tk.BOTH)
                
                # Index du trou
                hole_index = row * 3 + col
                self.game_state['hole_labels'].append(hole_label)
                
                # Binding pour le clic sur le trou - utilise une méthode de classe
                hole_label.bind("<Button-1>", lambda e, idx=hole_index: self.on_game_hole_click(idx))
        
        # Frame pour le message de résultat
        result_frame = tk.Frame(main_frame, bg="#1a1a2e")
        result_frame.pack(pady=(20, 10), fill=tk.X)
        
        self.game_state['result_label'] = tk.Label(
            result_frame,
            text="",
            font=("Arial", 11),
            bg="#1a1a2e",
            fg="#ffffff",
            wraplength=550,
            justify=tk.CENTER
        )
        self.game_state['result_label'].pack()
        
        # Frame pour les boutons
        button_frame = tk.Frame(main_frame, bg="#1a1a2e")
        button_frame.pack(pady=(15, 0))
        
        # Bouton rejouer (initialement caché)
        self.game_state['replay_btn'] = tk.Button(
            button_frame,
            text="🔄 Rejouer",
            command=self.restart_whack_game,
            bg="#059669",
            fg="white",
            font=("Arial", 12, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=10,
            activebackground="#047857",
            borderwidth=0
        )
        
        # Bouton fermer (initialement caché)
        self.game_state['close_btn'] = tk.Button(
            button_frame,
            text="✖ Fermer",
            command=self.close_whack_game,
            bg="#6b7280",
            fg="white",
            font=("Arial", 12, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=10,
            activebackground="#4b5563",
            borderwidth=0
        )
        
        # Effets hover pour les boutons
        self.game_state['replay_btn'].bind("<Enter>", lambda e: self.game_state['replay_btn'].config(bg="#047857"))
        self.game_state['replay_btn'].bind("<Leave>", lambda e: self.game_state['replay_btn'].config(bg="#059669"))
        self.game_state['close_btn'].bind("<Enter>", lambda e: self.game_state['close_btn'].config(bg="#4b5563"))
        self.game_state['close_btn'].bind("<Leave>", lambda e: self.game_state['close_btn'].config(bg="#6b7280"))
        
        # Fermer proprement si la fenêtre est fermée
        game_popup.protocol("WM_DELETE_WINDOW", self.close_whack_game)
        
        # Démarrer le jeu après un court délai
        game_popup.after(800, self.show_psyduck_random)
    
    def on_game_hole_click(self, hole_idx):
        """Gestionnaire de clic sur un trou du jeu"""
        if not self.game_state['game_active']:
            return
        
        if self.game_state['psyduck_visible'] and self.game_state['current_hole'] == hole_idx:
            # TOUCHÉ !
            self.game_state['hits'] += 1
            self.game_state['score_label'].config(text=f"🎯 Score: {self.game_state['hits']} / 3")
            
            # Annuler le timer de disparition automatique
            if self.game_state['game_timer']:
                self.game_state['popup'].after_cancel(self.game_state['game_timer'])
                self.game_state['game_timer'] = None
            
            # Afficher l'image psyduck-ko pendant 500ms pour simuler la frappe
            if self.game_state.get('psyduck_ko_image'):
                self.game_state['hole_labels'][hole_idx].config(
                    image=self.game_state['psyduck_ko_image'], 
                    bg="#ff6b6b"
                )
            else:
                # Fallback si l'image KO n'est pas chargée
                self.game_state['hole_labels'][hole_idx].config(bg="#ff6b6b")
            
            # Après 500ms, cacher l'image et continuer le jeu
            def after_ko_animation():
                self.game_state['hole_labels'][hole_idx].config(image='', text='', bg="#0d1117")
                self.game_state['psyduck_visible'] = False
                
                # Vérifier la victoire
                if self.game_state['hits'] >= self.game_state['target_hits']:
                    self.game_state['game_active'] = False
                    self.game_state['popup'].after(300, self.show_game_victory)
                else:
                    # Continuer le jeu après un court délai
                    self.game_state['popup'].after(300, self.show_psyduck_random)
            
            self.game_state['popup'].after(500, after_ko_animation)
    
    def show_psyduck_random(self):
        """Afficher le psyduck dans un trou aléatoire"""
        if not self.game_state['game_active']:
            return
        
        # Cacher l'ancien psyduck si visible
        self.hide_game_psyduck()
        
        # Choisir un nouveau trou aléatoire
        new_hole = random.randint(0, 8)
        self.game_state['current_hole'] = new_hole
        self.game_state['psyduck_visible'] = True
        
        # Afficher le psyduck
        if self.game_state['psyduck_image']:
            self.game_state['hole_labels'][new_hole].config(image=self.game_state['psyduck_image'], bg="#2d3748")
        else:
            self.game_state['hole_labels'][new_hole].config(text="🦆", font=("Arial", 40), bg="#2d3748")
        
        # Programmer la disparition après un délai (400-600ms)
        hide_delay = random.randint(400, 600)
        self.game_state['game_timer'] = self.game_state['popup'].after(hide_delay, self.auto_hide_psyduck)
    
    def auto_hide_psyduck(self):
        """Cacher le psyduck et en afficher un autre - RATÉ = score à 0"""
        if self.game_state['psyduck_visible'] and self.game_state['game_active']:
            # Le psyduck a disparu sans être cliqué = RATÉ
            if self.game_state['hits'] > 0:
                self.game_state['hits'] = 0
                self.game_state['score_label'].config(text="🎯 Score: 0 / 3", fg="#ff4444")
                self.game_state['popup'].after(300, lambda: self.game_state['score_label'].config(fg="#00ff88"))
            self.hide_game_psyduck()
            # Afficher un nouveau psyduck après un court délai
            self.game_state['popup'].after(300, self.show_psyduck_random)
    
    def hide_game_psyduck(self):
        """Cacher le psyduck actuel"""
        if self.game_state['current_hole'] is not None:
            self.game_state['hole_labels'][self.game_state['current_hole']].config(image='', text='', bg="#0d1117")
        self.game_state['psyduck_visible'] = False
        if self.game_state['game_timer']:
            self.game_state['popup'].after_cancel(self.game_state['game_timer'])
            self.game_state['game_timer'] = None
    
    def show_game_victory(self):
        """Afficher l'écran de victoire avec le message Easter Egg"""
        # Obtenir le message séquentiel
        message = self.show_easter_egg()
        
        # Afficher le message de victoire
        self.game_state['result_label'].config(
            text=f"🎉 BRAVO ! Tu as assommé Psykokwak !\n\n💬 \"{message}\"",
            fg="#ffd700",
            font=("Arial", 12, "bold")
        )
        
        # Forcer la mise à jour
        self.game_state['popup'].update_idletasks()
        
        # Afficher les boutons rejouer et fermer
        self.game_state['replay_btn'].pack(side=tk.LEFT, padx=10)
        self.game_state['close_btn'].pack(side=tk.LEFT, padx=10)
        
        # Forcer à nouveau la mise à jour pour s'assurer que les boutons sont visibles
        self.game_state['popup'].update_idletasks()
    
    def restart_whack_game(self):
        """Redémarrer le jeu"""
        self.game_state['hits'] = 0
        self.game_state['current_hole'] = None
        self.game_state['game_active'] = True
        self.game_state['psyduck_visible'] = False
        
        self.game_state['score_label'].config(text="🎯 Score: 0 / 3", fg="#00ff88")
        self.game_state['result_label'].config(text="", font=("Arial", 11))
        
        # Cacher les boutons
        self.game_state['replay_btn'].pack_forget()
        self.game_state['close_btn'].pack_forget()
        
        # Cacher tous les psyducks
        for label in self.game_state['hole_labels']:
            label.config(image='', text='', bg="#0d1117")
        
        # Redémarrer le jeu
        self.game_state['popup'].after(500, self.show_psyduck_random)
    
    def close_whack_game(self):
        """Fermer le jeu"""
        self.game_state['game_active'] = False
        if self.game_state['game_timer']:
            self.game_state['popup'].after_cancel(self.game_state['game_timer'])
        self.game_state['popup'].destroy()
    
    def show_casino_game(self):
        """Afficher le jeu de casino - Machine à sous Lautrec Game Corner ANIMÉ"""
        # Créer une fenêtre popup pour le casino
        casino_popup = tk.Toplevel(self.root)
        casino_popup.title("🎰 Lautrec Game Corner - Machine à Sous")
        casino_popup.geometry("650x850")
        casino_popup.configure(bg="#0a0a1e")
        casino_popup.resizable(False, False)
        
        # Centrer la fenêtre
        casino_popup.transient(self.root)
        casino_popup.grab_set()
        
        # Variables du casino (étendues pour les animations)
        self.casino_state = {
            'popup': casino_popup,
            'credits': getattr(self, 'saved_credits', 100),  # Charger les crédits sauvegardés
            'payout': 0,
            'spinning': False,
            'reel_labels': [],
            'reel_canvases': [],  # Canvas pour effets visuels
            'credit_label': None,
            'payout_label': None,
            'result_label': None,
            'spin_btn': None,
            'symbol_images': {},
            'reel_values': [0, 0, 0],
            'spin_timers': [None, None, None],
            'particles': [],  # Particules pour effets
            'glow_active': False,
            'shake_offset': [0, 0],
            'main_frame': None,
            'title_label': None,
            'animation_canvas': None,  # Canvas principal pour particules
            'reel_frames': [],
            'neon_pulse': 0,  # Pour effet néon pulsant
            # Nouveaux pour rouleaux verticaux
            'reel_strips': [],  # Liste des symboles pour chaque rouleau
            'reel_offsets': [0.0, 0.0, 0.0],  # Position de scroll verticale
            'reel_speeds': [0.0, 0.0, 0.0],  # Vitesse de scroll
            'reel_target_positions': [0, 0, 0],  # Position cible finale
        }
        
        # Liste des symboles du casino avec leurs valeurs
        self.casino_symbols = [
            {'name': '7', 'file': 'Celadon_Game_Corner_7_FRLG.png', 'value': 1000},
            {'name': 'Pikachu', 'file': 'Celadon_Game_Corner_Pikachu_FRLG.png', 'value': 150},
            {'name': 'Psyduck', 'file': 'Celadon_Game_Corner_Psyduck_FRLG.png', 'value': 130},
            {'name': 'Magnemite', 'file': 'Celadon_Game_Corner_Magnemite_FRLG.png', 'value': 128},
            {'name': 'Staryu', 'file': 'Celadon_Game_Corner_Staryu_FRLG.png', 'value': 125},
            {'name': 'Shellder', 'file': 'Celadon_Game_Corner_Shellder_FRLG.png', 'value': 123},
            {'name': 'Slowpoke', 'file': 'Celadon_Game_Corner_Slowpoke_FRLG.png', 'value': 122},
            {'name': 'Voltorb', 'file': 'Celadon_Game_Corner_Voltorb_FRLG.png', 'value': 120},
        ]
        
        # Charger les images des symboles
        for symbol in self.casino_symbols:
            try:
                symbol_path = os.path.join(os.path.dirname(__file__), "assets", symbol['file'])
                symbol_img = Image.open(symbol_path)
                symbol_img = symbol_img.resize((60, 60), Image.Resampling.LANCZOS)
                self.casino_state['symbol_images'][symbol['name']] = ImageTk.PhotoImage(symbol_img)
            except Exception as e:
                print(f"Erreur chargement symbole {symbol['name']}: {e}")
        
        # Créer les "strips" de symboles pour chaque rouleau (répétition pour loop infini)
        for i in range(3):
            # Mélanger les symboles pour chaque rouleau
            strip = list(range(len(self.casino_symbols))) * 3  # 3 répétitions
            random.shuffle(strip)
            self.casino_state['reel_strips'].append(strip)
        
        # Charger l'image de fond de la machine
        try:
            gamecorner_path = os.path.join(os.path.dirname(__file__), "assets", "gamecorner.png")
            gamecorner_img = Image.open(gamecorner_path)
            gamecorner_img = gamecorner_img.resize((450, 380), Image.Resampling.LANCZOS)
            self.casino_state['gamecorner_image'] = ImageTk.PhotoImage(gamecorner_img)
        except Exception as e:
            print(f"Erreur chargement gamecorner: {e}")
        
        # Canvas principal pour les effets de particules
        self.casino_state['animation_canvas'] = tk.Canvas(
            casino_popup,
            width=650,
            height=850,
            bg="#0a0a1e",
            highlightthickness=0
        )
        self.casino_state['animation_canvas'].place(x=0, y=0)
        
        # Frame principale (par-dessus le canvas)
        main_frame = tk.Frame(casino_popup, bg="#0a0a1e", padx=20, pady=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        self.casino_state['main_frame'] = main_frame
        
        # Titre avec effet néon pulsant
        title_label = tk.Label(
            main_frame,
            text="🎰 LAUTREC GAME CORNER 🎰",
            font=("Arial", 20, "bold"),
            bg="#0a0a1e",
            fg="#ffd700"
        )
        title_label.pack(pady=(0, 10))
        self.casino_state['title_label'] = title_label
        
        # Démarrer l'animation du titre néon
        self.animate_neon_title()
        
        # Frame pour afficher les crédits et payout avec style amélioré
        score_frame = tk.Frame(main_frame, bg="#1a1a3e", padx=20, pady=12, 
                              highlightbackground="#ffd700", highlightthickness=2)
        score_frame.pack(pady=(5, 15), fill=tk.X)
        
        credits_container = tk.Frame(score_frame, bg="#1a1a3e")
        credits_container.pack(side=tk.LEFT, expand=True)
        
        tk.Label(
            credits_container,
            text="💰 CRÉDITS",
            font=("Arial", 11, "bold"),
            bg="#1a1a3e",
            fg="#00ff88"
        ).pack()
        
        self.casino_state['credit_label'] = tk.Label(
            credits_container,
            text=str(self.casino_state['credits']),
            font=("Arial", 24, "bold"),
            bg="#1a1a3e",
            fg="#00ff88"
        )
        self.casino_state['credit_label'].pack()
        
        payout_container = tk.Frame(score_frame, bg="#1a1a3e")
        payout_container.pack(side=tk.RIGHT, expand=True)
        
        tk.Label(
            payout_container,
            text="🏆 GAIN",
            font=("Arial", 11, "bold"),
            bg="#1a1a3e",
            fg="#ffd700"
        ).pack()
        
        self.casino_state['payout_label'] = tk.Label(
            payout_container,
            text=str(self.casino_state['payout']),
            font=("Arial", 24, "bold"),
            bg="#1a1a3e",
            fg="#ffd700"
        )
        self.casino_state['payout_label'].pack()
        
        # Frame pour les rouleaux avec effets améliorés
        reels_outer_frame = tk.Frame(main_frame, bg="#8B4513", padx=12, pady=12,
                                     relief=tk.RIDGE, borderwidth=4)
        reels_outer_frame.pack(pady=15)
        
        # Sous-titre animé pour les rouleaux
        reel_title = tk.Label(
            reels_outer_frame,
            text="═══ ROULEAUX MAGIQUES ═══",
            font=("Courier", 10, "bold"),
            bg="#8B4513",
            fg="#ffd700"
        )
        reel_title.pack(pady=(0, 8))
        
        reels_frame = tk.Frame(reels_outer_frame, bg="#1a1a1a", padx=10, pady=10)
        reels_frame.pack()
        
        # Créer les 3 rouleaux avec Canvas pour animations VERTICALES
        for i in range(3):
            reel_container = tk.Frame(reels_frame, bg="#000000", 
                                     relief=tk.GROOVE, borderwidth=4)
            reel_container.pack(side=tk.LEFT, padx=8)
            self.casino_state['reel_frames'].append(reel_container)
            
            # Canvas pour le rouleau vertical (hauteur 3x pour afficher 3 symboles)
            reel_canvas = tk.Canvas(
                reel_container,
                width=90,
                height=210,  # 70px par symbole * 3
                bg="#ffffff",
                highlightthickness=0
            )
            reel_canvas.pack()
            self.casino_state['reel_canvases'].append(reel_canvas)
            
            # Dessiner les lignes de séparation entre symboles
            reel_canvas.create_line(0, 70, 90, 70, fill="#cccccc", width=1)
            reel_canvas.create_line(0, 140, 90, 140, fill="#cccccc", width=1)
            
            # Encadrer le symbole du milieu (celui qui compte)
            reel_canvas.create_rectangle(
                2, 72, 88, 138,
                outline="#ff0040",
                width=3
            )
            
            # Initialiser avec des symboles aléatoires
            initial_idx = random.randint(0, len(self.casino_symbols) - 1)
            self.casino_state['reel_values'][i] = initial_idx
            self.casino_state['reel_offsets'][i] = 0.0
            
        # Dessiner l'état initial des rouleaux
        self.draw_all_reels()
        
        # Label pour le résultat avec style amélioré
        self.casino_state['result_label'] = tk.Label(
            main_frame,
            text="✨ Appuyez sur SPIN pour jouer ! ✨\n(Coût: 10 crédits)",
            font=("Arial", 12, "bold"),
            bg="#0a0a1e",
            fg="#00ff88",
            wraplength=500
        )
        self.casino_state['result_label'].pack(pady=20)
        
        # Frame pour les boutons
        btn_frame = tk.Frame(main_frame, bg="#0a0a1e")
        btn_frame.pack(pady=15)
        
        # Bouton SPIN avec effet 3D et animation
        spin_btn_container = tk.Frame(btn_frame, bg="#0a0a1e")
        spin_btn_container.pack(side=tk.LEFT, padx=10)
        
        self.casino_state['spin_btn'] = tk.Button(
            spin_btn_container,
            text="🎰 SPIN !",
            command=self.casino_spin_animated,
            bg="#ff0040",
            fg="white",
            font=("Arial", 16, "bold"),
            cursor="hand2",
            relief=tk.RAISED,
            padx=50,
            pady=20,
            activebackground="#cc0030",
            borderwidth=5
        )
        self.casino_state['spin_btn'].pack()
        
        # Effet hover pour le bouton SPIN
        def spin_hover_enter(e):
            self.casino_state['spin_btn'].config(bg="#ff3366", relief=tk.RAISED, borderwidth=6)
        def spin_hover_leave(e):
            self.casino_state['spin_btn'].config(bg="#ff0040", relief=tk.RAISED, borderwidth=5)
        
        self.casino_state['spin_btn'].bind("<Enter>", spin_hover_enter)
        self.casino_state['spin_btn'].bind("<Leave>", spin_hover_leave)
        
        # Animation pulsante du bouton SPIN
        self.pulse_spin_button()
        
        # Bouton Fermer
        close_btn = tk.Button(
            btn_frame,
            text="✖ Quitter",
            command=self.close_casino_game,
            bg="#4b5563",
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=20,
            pady=12,
            activebackground="#374151",
            borderwidth=0
        )
        close_btn.pack(side=tk.LEFT, padx=10)
        
        # Bouton Inventaire des phrases
        inventory_btn = tk.Button(
            btn_frame,
            text="📜 Phrases",
            command=self.show_phrase_inventory,
            bg="#8b5cf6",
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=20,
            pady=12,
            activebackground="#7c3aed",
            borderwidth=0
        )
        inventory_btn.pack(side=tk.LEFT, padx=10)
        
        # Effets hover
        def inv_on_enter(e):
            inventory_btn.config(bg="#7c3aed")
        def inv_on_leave(e):
            inventory_btn.config(bg="#8b5cf6")
        inventory_btn.bind("<Enter>", inv_on_enter)
        inventory_btn.bind("<Leave>", inv_on_leave)
        
        def close_on_enter(e):
            close_btn.config(bg="#374151")
        def close_on_leave(e):
            close_btn.config(bg="#4b5563")
        close_btn.bind("<Enter>", close_on_enter)
        close_btn.bind("<Leave>", close_on_leave)
        
        # Compteur de phrases débloquées
        total_phrases = len(self.easter_messages)
        unlocked_phrases = len(self.purchased_phrases)
        
        phrases_label = tk.Label(
            main_frame,
            text=f"📜 Phrases débloquées: {unlocked_phrases}/{total_phrases} (35 crédits/phrase)",
            font=("Arial", 10),
            bg="#0a0a1e",
            fg="#a78bfa"
        )
        phrases_label.pack(pady=(5, 10))
        self.casino_state['phrases_label'] = phrases_label
        
        # Table des gains avec style amélioré
        paytable_frame = tk.LabelFrame(
            main_frame,
            text="🏆 TABLE DES GAINS 🏆",
            font=("Arial", 11, "bold"),
            bg="#0a0a1e",
            fg="#ffd700",
            padx=15,
            pady=8,
            highlightbackground="#ffd700",
            highlightthickness=1
        )
        paytable_frame.pack(pady=10, fill=tk.X)
        
        paytable_text = "777 = 1000 💰 | Pikachu x3 = 500 💰 | Psyduck x3 = 300 💰\n2 identiques = valeur/3"
        tk.Label(
            paytable_frame,
            text=paytable_text,
            font=("Arial", 10, "bold"),
            bg="#0a0a1e",
            fg="#00ff88",
            wraplength=550
        ).pack()
        
        # Fermer proprement si la fenêtre est fermée
        casino_popup.protocol("WM_DELETE_WINDOW", self.close_casino_game)
    
    def draw_all_reels(self):
        """Dessiner tous les rouleaux avec leurs symboles visibles"""
        for i in range(3):
            self.draw_reel(i)
    
    def draw_reel(self, reel_idx):
        """Dessiner un rouleau avec 3 symboles visibles (effet de défilement vertical fluide)"""
        canvas = self.casino_state['reel_canvases'][reel_idx]
        strip = self.casino_state['reel_strips'][reel_idx]
        offset = self.casino_state['reel_offsets'][reel_idx]
        
        # Nettoyer le canvas (sauf les lignes et le cadre)
        canvas.delete("symbol")
        
        # Position de base dans le strip
        strip_len = len(strip)
        # Le symbole à l'index (offset) sera au centre (ligne du milieu = slot 1)
        base_pos = int(offset) % strip_len
        
        # Offset pixel pour le défilement fluide
        pixel_offset = (offset % 1.0) * 70
        
        # Dessiner 5 symboles pour assurer une couverture complète pendant le scroll
        for slot in range(-1, 4):
            # Position dans le strip (slot 1 = centre = base_pos)
            strip_pos = (base_pos + slot - 1) % strip_len
            symbol_idx = strip[strip_pos]
            symbol = self.casino_symbols[symbol_idx]
            
            # Position Y : slot 1 (centre) est à 70px, on soustrait le pixel_offset pour défiler vers le haut
            y_center = (slot * 70) + 35 - pixel_offset
            
            # Dessiner le symbole s'il est dans la zone visible
            if -35 <= y_center <= 245:
                if symbol['name'] in self.casino_state['symbol_images']:
                    img = self.casino_state['symbol_images'][symbol['name']]
                    canvas.create_image(
                        45, y_center,
                        image=img,
                        tags="symbol"
                    )
        
        # Redessiner le cadre rouge du milieu (au-dessus des symboles)
        canvas.delete("frame")
        canvas.create_rectangle(
            2, 72, 88, 138,
            outline="#ff0040",
            width=3,
            tags="frame"
        )
    
    def animate_neon_title(self):
        """Animation de pulsation néon pour le titre"""
        if not hasattr(self, 'casino_state') or 'title_label' not in self.casino_state:
            return
        
        try:
            # Calculer la couleur avec effet de pulsation
            import math
            self.casino_state['neon_pulse'] += 0.15
            brightness = int(200 + 55 * math.sin(self.casino_state['neon_pulse']))
            color = f'#{brightness:02x}d700'
            
            self.casino_state['title_label'].config(fg=color)
            
            # Continuer l'animation
            if self.casino_state.get('popup') and self.casino_state['popup'].winfo_exists():
                self.casino_state['popup'].after(50, self.animate_neon_title)
        except:
            pass
    
    def pulse_spin_button(self):
        """Animation de pulsation du bouton SPIN"""
        if not hasattr(self, 'casino_state') or 'spin_btn' not in self.casino_state:
            return
        
        try:
            if not self.casino_state.get('spinning', False):
                # Effet de pulsation quand le bouton est disponible
                import math
                pulse = math.sin(self.casino_state.get('neon_pulse', 0) * 1.5)
                scale = 1.0 + pulse * 0.02
                
                # Alterner légèrement la taille via le padding
                if pulse > 0:
                    self.casino_state['spin_btn'].config(padx=52, pady=21)
                else:
                    self.casino_state['spin_btn'].config(padx=50, pady=20)
            
            # Continuer l'animation
            if self.casino_state.get('popup') and self.casino_state['popup'].winfo_exists():
                self.casino_state['popup'].after(100, self.pulse_spin_button)
        except:
            pass
    
    def create_particle(self, x, y, color, size=4, vx=None, vy=None):
        """Créer une particule pour les effets visuels"""
        import math
        import random as rand
        
        if vx is None:
            vx = rand.uniform(-3, 3)
        if vy is None:
            vy = rand.uniform(-5, -2)
        
        particle = {
            'x': x,
            'y': y,
            'vx': vx,
            'vy': vy,
            'color': color,
            'size': size,
            'life': 50,
            'id': None
        }
        
        # Dessiner la particule
        canvas = self.casino_state['animation_canvas']
        particle['id'] = canvas.create_oval(
            x - size, y - size, x + size, y + size,
            fill=color, outline=color
        )
        
        self.casino_state['particles'].append(particle)
    
    def update_particles(self):
        """Mettre à jour et animer les particules"""
        if not hasattr(self, 'casino_state') or 'particles' not in self.casino_state:
            return
        
        try:
            canvas = self.casino_state['animation_canvas']
            particles_to_remove = []
            
            for particle in self.casino_state['particles']:
                # Mise à jour physique
                particle['vy'] += 0.2  # Gravité
                particle['x'] += particle['vx']
                particle['y'] += particle['vy']
                particle['life'] -= 1
                
                # Mettre à jour la position
                if particle['id']:
                    canvas.coords(
                        particle['id'],
                        particle['x'] - particle['size'],
                        particle['y'] - particle['size'],
                        particle['x'] + particle['size'],
                        particle['y'] + particle['size']
                    )
                    
                    # Fading
                    alpha = particle['life'] / 50.0
                    if alpha < 0.3:
                        canvas.itemconfig(particle['id'], state='hidden')
                
                # Marquer pour suppression si morte
                if particle['life'] <= 0:
                    particles_to_remove.append(particle)
            
            # Nettoyer les particules mortes
            for particle in particles_to_remove:
                if particle['id']:
                    canvas.delete(particle['id'])
                self.casino_state['particles'].remove(particle)
            
            # Continuer l'animation
            if self.casino_state.get('popup') and self.casino_state['popup'].winfo_exists():
                self.casino_state['popup'].after(30, self.update_particles)
        except:
            pass
    
    def shake_screen(self, intensity=10, duration=20):
        """Effet de tremblement de l'écran"""
        import random as rand
        
        def _shake(count):
            if count <= 0 or not self.casino_state.get('popup'):
                # Remettre en position normale
                self.casino_state['main_frame'].place(x=0, y=0)
                return
            
            # Déplacement aléatoire
            offset_x = rand.randint(-intensity, intensity)
            offset_y = rand.randint(-intensity, intensity)
            
            # Diminuer l'intensité progressivement
            current_intensity = int(intensity * (count / duration))
            offset_x = rand.randint(-current_intensity, current_intensity)
            offset_y = rand.randint(-current_intensity, current_intensity)
            
            self.casino_state['main_frame'].place(x=offset_x, y=offset_y)
            
            # Continuer le shake
            self.casino_state['popup'].after(30, lambda: _shake(count - 1))
        
        _shake(duration)
    
    def flash_screen(self, color="#ffffff", times=3):
        """Effet de flash d'écran"""
        def _flash(count):
            if count <= 0 or not self.casino_state.get('popup'):
                self.casino_state['animation_canvas'].config(bg="#0a0a1e")
                return
            
            # Alterner entre la couleur et le fond normal
            if count % 2 == 0:
                self.casino_state['animation_canvas'].config(bg=color)
            else:
                self.casino_state['animation_canvas'].config(bg="#0a0a1e")
            
            self.casino_state['popup'].after(80, lambda: _flash(count - 1))
        
        _flash(times * 2)
    
    def animate_counter(self, label, start, end, duration=500, prefix="", color_start="#ffffff", color_end=None):
        """Animer un compteur de nombre"""
        if color_end is None:
            color_end = color_start
        
        import math
        steps = 20
        delay = duration // steps
        
        def _update(step):
            if step > steps or not self.casino_state.get('popup'):
                label.config(text=f"{prefix}{end}", fg=color_end)
                return
            
            # Easing function (ease-out)
            progress = step / steps
            eased = 1 - math.pow(1 - progress, 3)
            current = int(start + (end - start) * eased)
            
            # Interpolation de couleur
            if color_start != color_end and step % 2 == 0:
                label.config(fg=color_end if step > steps // 2 else color_start)
            
            label.config(text=f"{prefix}{current}")
            self.casino_state['popup'].after(delay, lambda: _update(step + 1))
        
        _update(0)
    
    def create_win_particles(self, x, y, count=30):
        """Créer des particules de victoire"""
        import random as rand
        colors = ["#ffd700", "#ffed4e", "#ff0", "#fff", "#00ff88"]
        
        for _ in range(count):
            angle = rand.uniform(0, 6.28)  # 2*pi
            speed = rand.uniform(3, 8)
            vx = speed * math.cos(angle)
            vy = speed * math.sin(angle) - rand.uniform(2, 4)
            size = rand.randint(3, 7)
            color = rand.choice(colors)
            
            self.create_particle(x, y, color, size, vx, vy)
    
    def glow_reel(self, reel_idx, color="#ffd700"):
        """Effet de lueur sur un rouleau gagnant (version verticale)"""
        canvas = self.casino_state['reel_canvases'][reel_idx]
        
        def _glow(step):
            if step <= 0 or not self.casino_state.get('popup'):
                canvas.config(bg="#ffffff")
                # Redessiner le cadre rouge
                canvas.delete("frame")
                canvas.create_rectangle(
                    2, 72, 88, 138,
                    outline="#ff0040",
                    width=3,
                    tags="frame"
                )
                return
            
            # Alterner entre la couleur glow et blanc
            if step % 2 == 0:
                canvas.config(bg=color)
                # Cadre de la couleur glow aussi
                canvas.delete("frame")
                canvas.create_rectangle(
                    2, 72, 88, 138,
                    outline=color,
                    width=4,
                    tags="frame"
                )
            else:
                canvas.config(bg="#ffffff")
                canvas.delete("frame")
                canvas.create_rectangle(
                    2, 72, 88, 138,
                    outline="#ff0040",
                    width=3,
                    tags="frame"
                )
            
            self.casino_state['popup'].after(100, lambda: _glow(step - 1))
        
        _glow(8)
    
    def casino_spin_animated(self):
        """Version animée de casino_spin avec effets"""
        # Effet de pression sur le bouton
        btn = self.casino_state['spin_btn']
        btn.config(relief=tk.SUNKEN, borderwidth=3)
        self.casino_state['popup'].after(100, lambda: btn.config(relief=tk.RAISED, borderwidth=5))
        
        # Flash rapide
        self.flash_screen("#ff0040", times=1)
        
        # Lancer le spin normal après l'effet
        self.casino_state['popup'].after(150, self.casino_spin)
    
    def casino_spin(self):
        """Lancer les rouleaux du casino avec défilement vertical"""
        if self.casino_state['spinning']:
            return
        
        # Vérifier les crédits
        if self.casino_state['credits'] < 10:
            self.casino_state['result_label'].config(
                text="❌ Pas assez de crédits ! Minimum 10 crédits requis.",
                fg="#ff4444"
            )
            self.play_sound('error')
            return
        
        # Jouer le son de spin
        self.play_sound('spin')
        
        # Déduire les crédits avec animation
        old_credits = self.casino_state['credits']
        self.casino_state['credits'] -= 10
        self.animate_counter(
            self.casino_state['credit_label'],
            old_credits,
            self.casino_state['credits'],
            duration=300,
            color_start="#00ff88",
            color_end="#00ff88"
        )
        
        # Sauvegarder les crédits immédiatement après déduction
        self.save_purchased_phrases()
        
        self.casino_state['payout'] = 0
        self.casino_state['payout_label'].config(text="0")
        self.casino_state['result_label'].config(
            text="🎰 Les rouleaux tournent...",
            fg="#ffffff",
            font=("Arial", 11, "bold")
        )
        self.casino_state['spinning'] = True
        self.casino_state['spin_btn'].config(state=tk.DISABLED)
        
        # Déterminer les symboles finaux pour chaque rouleau
        final_symbols = [
            random.randint(0, len(self.casino_symbols) - 1),
            random.randint(0, len(self.casino_symbols) - 1),
            random.randint(0, len(self.casino_symbols) - 1)
        ]
        self.casino_state['reel_values'] = final_symbols
        
        # Configuration du spin pour chaque rouleau
        # Chaque rouleau tourne différemment et s'arrête à des moments différents
        base_rotations = 20  # Nombre de symboles à faire défiler
        
        for i in range(3):
            # Nombre de rotations pour ce rouleau (augmente pour chaque rouleau)
            rotations = base_rotations + (i * 8)
            
            # Position cible finale (alignée sur le symbole gagnant dans le strip)
            # On cherche où se trouve le symbole final dans notre strip
            target_symbol = final_symbols[i]
            strip = self.casino_state['reel_strips'][i]
            strip_len = len(strip)
            
            # Trouver la première occurrence du symbole cible après avoir fait plusieurs rotations
            current_pos = int(self.casino_state['reel_offsets'][i])
            # On veut que le symbole soit visible après 'rotations' symboles défilés
            search_start = current_pos + rotations
            
            for j in range(strip_len):
                check_pos = (search_start + j) % strip_len
                if strip[check_pos] == target_symbol:
                    # La position cible est search_start + j (en prenant en compte le cycle)
                    target_offset = search_start + j
                    break
            
            self.casino_state['reel_target_positions'][i] = target_offset
            self.casino_state['reel_speeds'][i] = 8.0  # Vitesse initiale
        
        # Démarrer l'animation des rouleaux
        self.animate_vertical_reels()
    
    def animate_vertical_reels(self):
        """Animer les rouleaux avec défilement vertical fluide et ralentissement progressif"""
        if not self.casino_state.get('spinning'):
            return
        
        all_stopped = True
        
        for i in range(3):
            current_offset = self.casino_state['reel_offsets'][i]
            target_offset = self.casino_state['reel_target_positions'][i]
            speed = self.casino_state['reel_speeds'][i]
            
            # Distance restante
            distance = target_offset - current_offset
            
            if distance > 0.05:  # Pas encore arrêté (seuil réduit pour précision)
                all_stopped = False
                
                # Ralentissement progressif avec easing fluide
                if distance < 3:
                    # Phase finale : très lent pour bien voir l'alignement
                    speed = max(0.08, distance * 0.15)
                elif distance < 8:
                    # Phase de ralentissement prononcé
                    speed = max(0.3, distance * 0.2)
                elif distance < 15:
                    # Phase de décélération
                    speed = max(1.5, speed * 0.92)
                else:
                    # Phase de vitesse rapide
                    speed = max(5.0, speed * 0.98)
                
                # Mettre à jour la position
                self.casino_state['reel_offsets'][i] += speed
                self.casino_state['reel_speeds'][i] = speed
                
                # S'assurer de ne pas dépasser la cible
                if self.casino_state['reel_offsets'][i] >= target_offset:
                    self.casino_state['reel_offsets'][i] = target_offset
                    # Son d'arrêt du rouleau
                    self.play_sound('reel_stop')
                
                # Redessiner le rouleau
                self.draw_reel(i)
            else:
                # Ce rouleau est arrêté, s'assurer qu'il est exactement sur la position entière
                self.casino_state['reel_offsets'][i] = float(int(target_offset))
                self.draw_reel(i)
        
        # Continuer l'animation ou vérifier le résultat
        if not all_stopped:
            self.casino_state['popup'].after(16, self.animate_vertical_reels)  # ~60 FPS
        else:
            # Tous les rouleaux sont arrêtés - redessiner une dernière fois pour s'assurer de l'alignement
            for i in range(3):
                self.casino_state['reel_offsets'][i] = float(int(self.casino_state['reel_target_positions'][i]))
                self.draw_reel(i)
            self.casino_state['popup'].after(300, self.check_casino_result)
    
    def check_casino_result(self):
        """Vérifier le résultat et calculer les gains AVEC ANIMATIONS"""
        import math
        
        self.casino_state['spinning'] = False
        self.casino_state['spin_btn'].config(state=tk.NORMAL)
        
        # Obtenir les symboles finaux
        symbols = [self.casino_symbols[idx] for idx in self.casino_state['reel_values']]
        names = [s['name'] for s in symbols]
        
        payout = 0
        result_text = ""
        win_type = None  # 'jackpot', 'big_win', 'win', 'near_miss', 'lose'
        
        # Vérifier les combinaisons gagnantes
        if names[0] == names[1] == names[2]:
            # 3 symboles identiques - JACKPOT!
            payout = symbols[0]['value']
            if names[0] == '7':
                result_text = f"💎 MEGA JACKPOT! 💎\nTriple 7! +{payout} crédits!"
                win_type = 'jackpot'
            else:
                result_text = f"✨ SUPER VICTOIRE! ✨\n3x {names[0]}! +{payout} crédits!"
                win_type = 'big_win' if payout >= 100 else 'win'
        elif names[0] == names[1] or names[1] == names[2] or names[0] == names[2]:
            # 2 symboles identiques
            if names[0] == names[1]:
                payout = symbols[0]['value'] // 3
                matching_reels = [0, 1]
            elif names[1] == names[2]:
                payout = symbols[1]['value'] // 3
                matching_reels = [1, 2]
            else:
                payout = symbols[0]['value'] // 3
                matching_reels = [0, 2]
            payout = max(payout, 2)
            result_text = f"👍 Pas mal! 2 identiques!\n+{payout} crédits!"
            win_type = 'win'
        else:
            # Vérifier le "near miss" (2 symboles identiques parmi les 3 qui se suivent)
            result_text = "😔 Dommage... Réessayez!"
            win_type = 'lose'
        
        # ANIMATIONS SELON LE TYPE DE VICTOIRE
        if win_type == 'jackpot':
            # JACKPOT - Animations maximales!
            self.play_sound('jackpot')
            
            # 1. Shake violent de l'écran
            self.shake_screen(intensity=15, duration=30)
            
            # 2. Flash d'écran doré
            self.flash_screen("#ffd700", times=5)
            
            # 3. Effet glow sur TOUS les rouleaux
            for i in range(3):
                self.glow_reel(i, "#ffd700")
            
            # 4. Explosion de particules dorées
            self.casino_state['popup'].after(200, lambda: self.create_win_particles(325, 400, count=80))
            self.casino_state['popup'].after(400, lambda: self.create_win_particles(325, 400, count=60))
            
            # 5. Démarrer l'animation des particules
            self.update_particles()
            
            # 6. Animation du texte de résultat (zoom effet)
            self.animate_result_text(result_text, "#ffd700", scale_effect=True)
            
        elif win_type == 'big_win':
            # Grande victoire - Animations importantes
            self.play_sound('jackpot')
            
            # Shake modéré
            self.shake_screen(intensity=10, duration=20)
            
            # Flash
            self.flash_screen("#00ff88", times=3)
            
            # Glow sur tous les rouleaux
            for i in range(3):
                self.glow_reel(i, "#00ff88")
            
            # Particules vertes
            self.casino_state['popup'].after(200, lambda: self.create_win_particles(325, 400, count=50))
            self.update_particles()
            
            self.animate_result_text(result_text, "#00ff88", scale_effect=True)
            
        elif win_type == 'win':
            # Victoire normale - Animations légères
            self.play_sound('win')
            
            # Flash léger
            self.flash_screen("#00ff88", times=2)
            
            # Glow sur les rouleaux gagnants si définis
            if 'matching_reels' in locals():
                for reel_idx in matching_reels:
                    self.glow_reel(reel_idx, "#00ff88")
            
            # Quelques particules
            self.casino_state['popup'].after(150, lambda: self.create_win_particles(325, 400, count=25))
            self.update_particles()
            
            self.animate_result_text(result_text, "#00ff88")
            
        else:
            # Perte - Pas d'animations, juste le son
            self.play_sound('lose')
            self.casino_state['result_label'].config(text=result_text, fg="#ff6b6b")
        
        # Mettre à jour les gains avec animation
        if payout > 0:
            old_credits = self.casino_state['credits']
            new_credits = old_credits + payout
            
            # Animer le compteur de crédits
            self.animate_counter(
                self.casino_state['credit_label'],
                old_credits,
                new_credits,
                duration=800,
                color_start="#00ff88",
                color_end="#00ff88"
            )
            
            # Animer le payout
            self.animate_counter(
                self.casino_state['payout_label'],
                0,
                payout,
                duration=600,
                color_start="#ffd700",
                color_end="#ffd700"
            )
            
            self.casino_state['credits'] = new_credits
            self.casino_state['payout'] = payout
            
            # Sauvegarder les crédits automatiquement après chaque gain
            self.save_purchased_phrases()
        
        # Mettre à jour le compteur de phrases
        if 'phrases_label' in self.casino_state:
            total_phrases = len(self.easter_messages)
            unlocked_phrases = len(self.purchased_phrases)
            self.casino_state['phrases_label'].config(
                text=f"📜 Phrases débloquées: {unlocked_phrases}/{total_phrases} (35 crédits/phrase)"
            )
        
        # Vérifier si le joueur est ruiné
        if self.casino_state['credits'] < 3:
            bonus_text = result_text + "\n\n💸 Plus assez de crédits!\n🎁 BONUS: +50 crédits!"
            self.casino_state['result_label'].config(
                text=bonus_text,
                fg="#ffd700"
            )
            old_credits = self.casino_state['credits']
            self.casino_state['credits'] += 50
            
            # Sauvegarder le bonus
            self.save_purchased_phrases()
            
            # Animer le bonus
            self.casino_state['popup'].after(1500, lambda: self.animate_counter(
                self.casino_state['credit_label'],
                old_credits,
                self.casino_state['credits'],
                duration=500,
                color_start="#ffd700",
                color_end="#00ff88"
            ))
            
            # Flash pour le bonus
            self.casino_state['popup'].after(1500, lambda: self.flash_screen("#ffd700", times=2))
    
    def animate_result_text(self, text, color, scale_effect=False):
        """Animer l'apparition du texte de résultat"""
        label = self.casino_state['result_label']
        
        if scale_effect:
            # Effet de zoom pour grandes victoires
            sizes = [8, 10, 12, 14, 12, 11]
            def _zoom(step):
                if step >= len(sizes):
                    label.config(font=("Arial", 12, "bold"), fg=color, text=text)
                    return
                label.config(font=("Arial", sizes[step], "bold"), fg=color, text=text)
                self.casino_state['popup'].after(60, lambda: _zoom(step + 1))
            _zoom(0)
        else:
            # Apparition simple
            label.config(text=text, fg=color, font=("Arial", 11, "bold"))
    
    def close_casino_game(self):
        """Fermer le jeu de casino et sauvegarder les crédits"""
        # Sauvegarder les crédits avant de fermer
        self.save_purchased_phrases()
        
        # Annuler tous les timers en cours
        for timer in self.casino_state.get('spin_timers', []):
            if timer:
                try:
                    self.casino_state['popup'].after_cancel(timer)
                except:
                    pass
        self.casino_state['popup'].destroy()
    
    def show_phrase_inventory(self):
        """Afficher l'inventaire des phrases à acheter"""
        # Créer une fenêtre popup pour l'inventaire
        inventory_popup = tk.Toplevel(self.casino_state['popup'])
        inventory_popup.title("📜 Inventaire des Phrases Secrètes")
        inventory_popup.geometry("700x600")
        inventory_popup.configure(bg="#1a1a2e")
        inventory_popup.resizable(False, False)
        
        # Centrer la fenêtre
        inventory_popup.transient(self.casino_state['popup'])
        inventory_popup.grab_set()
        
        # Frame principale
        main_frame = tk.Frame(inventory_popup, bg="#1a1a2e", padx=20, pady=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Titre
        title_label = tk.Label(
            main_frame,
            text="📜 PHRASES SECRÈTES 📜",
            font=("Arial", 16, "bold"),
            bg="#1a1a2e",
            fg="#ffd700"
        )
        title_label.pack(pady=(0, 5))
        
        # Sous-titre avec compteur
        total_phrases = len(self.easter_messages)
        unlocked_count = len(self.purchased_phrases)
        
        subtitle_label = tk.Label(
            main_frame,
            text=f"Débloquées: {unlocked_count}/{total_phrases} | Coût: 35 crédits par phrase",
            font=("Arial", 10),
            bg="#1a1a2e",
            fg="#a78bfa"
        )
        subtitle_label.pack(pady=(0, 15))
        
        # Frame avec scrollbar pour les phrases
        canvas_frame = tk.Frame(main_frame, bg="#1a1a2e")
        canvas_frame.pack(fill=tk.BOTH, expand=True)
        
        canvas = tk.Canvas(canvas_frame, bg="#1a1a2e", highlightthickness=0)
        scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg="#1a1a2e")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Permettre le scroll avec la molette
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Stocker les références pour mise à jour
        self.inventory_state = {
            'popup': inventory_popup,
            'phrase_labels': [],
            'buy_buttons': [],
            'subtitle_label': subtitle_label,
            'scrollable_frame': scrollable_frame
        }
        
        # Afficher chaque phrase
        for idx, phrase in enumerate(self.easter_messages):
            phrase_frame = tk.Frame(scrollable_frame, bg="#16213e", padx=10, pady=8)
            phrase_frame.pack(fill=tk.X, pady=3, padx=5)
            
            # Numéro de la phrase
            num_label = tk.Label(
                phrase_frame,
                text=f"#{idx + 1}",
                font=("Arial", 9, "bold"),
                bg="#16213e",
                fg="#6b7280",
                width=4
            )
            num_label.pack(side=tk.LEFT, padx=(0, 10))
            
            # Texte de la phrase (masqué ou visible)
            is_purchased = idx in self.purchased_phrases
            
            if is_purchased:
                display_text = phrase
                text_color = "#ffffff"
            else:
                # Masquer le texte
                display_text = "🔒 " + "•" * min(len(phrase) // 2, 40) + " [PHRASE SECRÈTE]"
                text_color = "#6b7280"
            
            phrase_label = tk.Label(
                phrase_frame,
                text=display_text,
                font=("Arial", 9),
                bg="#16213e",
                fg=text_color,
                wraplength=450,
                justify=tk.LEFT,
                anchor="w"
            )
            phrase_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            # Bouton d'achat (si non acheté)
            if not is_purchased:
                buy_btn = tk.Button(
                    phrase_frame,
                    text="🔓 35💰",
                    command=lambda i=idx: self.purchase_phrase(i),
                    bg="#059669",
                    fg="white",
                    font=("Arial", 9, "bold"),
                    cursor="hand2",
                    relief=tk.FLAT,
                    padx=10,
                    pady=3,
                    activebackground="#047857",
                    borderwidth=0
                )
                buy_btn.pack(side=tk.RIGHT, padx=(10, 0))
                
                # Hover effect
                buy_btn.bind("<Enter>", lambda e, b=buy_btn: b.config(bg="#047857"))
                buy_btn.bind("<Leave>", lambda e, b=buy_btn: b.config(bg="#059669"))
                
                self.inventory_state['buy_buttons'].append((idx, buy_btn, phrase_frame, phrase_label))
            else:
                # Badge "Débloqué"
                unlocked_badge = tk.Label(
                    phrase_frame,
                    text="✅",
                    font=("Arial", 10),
                    bg="#16213e",
                    fg="#00ff88"
                )
                unlocked_badge.pack(side=tk.RIGHT, padx=(10, 0))
        
        # Bouton Fermer
        close_btn = tk.Button(
            main_frame,
            text="✖ Fermer",
            command=lambda: self.close_inventory(inventory_popup),
            bg="#6b7280",
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=30,
            pady=10,
            activebackground="#4b5563",
            borderwidth=0
        )
        close_btn.pack(pady=(15, 0))
        
        # Unbind mousewheel when closing
        inventory_popup.protocol("WM_DELETE_WINDOW", lambda: self.close_inventory(inventory_popup))
    
    def close_inventory(self, popup):
        """Fermer l'inventaire et nettoyer les bindings"""
        try:
            popup.unbind_all("<MouseWheel>")
        except:
            pass
        popup.destroy()
    
    def purchase_phrase(self, phrase_idx):
        """Acheter une phrase avec les crédits du casino"""
        cost = 35
        
        # Vérifier les crédits
        if self.casino_state['credits'] < cost:
            self.play_sound('error')
            messagebox.showwarning(
                "Crédits insuffisants",
                f"Vous avez besoin de {cost} crédits pour acheter cette phrase.\n"
                f"Crédits actuels: {self.casino_state['credits']}"
            )
            return
        
        # Déduire les crédits
        self.casino_state['credits'] -= cost
        self.casino_state['credit_label'].config(text=str(self.casino_state['credits']))
        
        # Marquer la phrase comme achetée
        self.purchased_phrases.add(phrase_idx)
        self.save_purchased_phrases()
        
        # Jouer le son d'achat
        self.play_sound('purchase')
        
        # Mettre à jour l'affichage dans l'inventaire
        phrase_text = self.easter_messages[phrase_idx]
        
        # Trouver et mettre à jour le bouton/label correspondant
        for idx, buy_btn, phrase_frame, phrase_label in self.inventory_state['buy_buttons']:
            if idx == phrase_idx:
                # Supprimer le bouton d'achat
                buy_btn.destroy()
                
                # Mettre à jour le texte de la phrase
                phrase_label.config(
                    text=phrase_text,
                    fg="#ffffff"
                )
                
                # Ajouter le badge débloqué
                unlocked_badge = tk.Label(
                    phrase_frame,
                    text="✅",
                    font=("Arial", 10),
                    bg="#16213e",
                    fg="#00ff88"
                )
                unlocked_badge.pack(side=tk.RIGHT, padx=(10, 0))
                break
        
        # Mettre à jour le compteur
        total_phrases = len(self.easter_messages)
        unlocked_count = len(self.purchased_phrases)
        self.inventory_state['subtitle_label'].config(
            text=f"Débloquées: {unlocked_count}/{total_phrases} | Coût: 35 crédits par phrase"
        )
        
        # Mettre à jour le label dans le casino
        if 'phrases_label' in self.casino_state:
            self.casino_state['phrases_label'].config(
                text=f"📜 Phrases débloquées: {unlocked_count}/{total_phrases} (35 crédits/phrase)"
            )
    
    def show_add_phrases_window(self):
        """Afficher la fenêtre pour ajouter des nouvelles phrases au shop"""
        # Créer une fenêtre popup
        add_popup = tk.Toplevel(self.root)
        add_popup.title("📝 Ajouter des Phrases Secrètes")
        add_popup.geometry("750x700")
        add_popup.configure(bg="#1a1a2e")
        add_popup.resizable(False, False)
        
        # Centrer la fenêtre
        add_popup.transient(self.root)
        add_popup.grab_set()
        
        # Variables pour stocker les nouvelles phrases
        self.new_phrases_list = []
        
        # Frame principale
        main_frame = tk.Frame(add_popup, bg="#1a1a2e", padx=20, pady=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Titre
        title_label = tk.Label(
            main_frame,
            text="📝 AJOUTER DES PHRASES SECRÈTES 📝",
            font=("Arial", 16, "bold"),
            bg="#1a1a2e",
            fg="#ffd700"
        )
        title_label.pack(pady=(0, 5))
        
        # Sous-titre explicatif
        subtitle_label = tk.Label(
            main_frame,
            text="Ajoutez vos propres phrases secrètes qui seront cryptées et ajoutées au shop !",
            font=("Arial", 10),
            bg="#1a1a2e",
            fg="#a78bfa"
        )
        subtitle_label.pack(pady=(0, 15))
        
        # Frame pour la saisie de nouvelle phrase
        input_frame = tk.LabelFrame(
            main_frame,
            text="✏️ Nouvelle phrase",
            font=("Arial", 11, "bold"),
            bg="#1a1a2e",
            fg="#00ff88",
            padx=15,
            pady=10
        )
        input_frame.pack(fill=tk.X, pady=(0, 15))
        
        # Zone de texte pour saisir la phrase
        self.new_phrase_entry = tk.Text(
            input_frame,
            height=3,
            font=("Arial", 11),
            bg="#16213e",
            fg="#ffffff",
            insertbackground="#ffffff",
            wrap=tk.WORD,
            relief=tk.FLAT,
            padx=10,
            pady=10
        )
        self.new_phrase_entry.pack(fill=tk.X, pady=(5, 10))
        
        # Bouton pour ajouter la phrase à la liste temporaire
        add_btn = tk.Button(
            input_frame,
            text="➕ Ajouter à la liste",
            command=lambda: self.add_phrase_to_temp_list(add_popup),
            bg="#059669",
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=20,
            pady=8,
            activebackground="#047857",
            borderwidth=0
        )
        add_btn.pack()
        
        # Hover effect
        add_btn.bind("<Enter>", lambda e: add_btn.config(bg="#047857"))
        add_btn.bind("<Leave>", lambda e: add_btn.config(bg="#059669"))
        
        # Frame pour afficher les phrases en attente
        pending_frame = tk.LabelFrame(
            main_frame,
            text="📋 Phrases en attente d'ajout au shop",
            font=("Arial", 11, "bold"),
            bg="#1a1a2e",
            fg="#ffd700",
            padx=15,
            pady=10
        )
        pending_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        # Canvas avec scrollbar pour les phrases en attente
        canvas_frame = tk.Frame(pending_frame, bg="#1a1a2e")
        canvas_frame.pack(fill=tk.BOTH, expand=True)
        
        pending_canvas = tk.Canvas(canvas_frame, bg="#1a1a2e", highlightthickness=0)
        pending_scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=pending_canvas.yview)
        self.pending_phrases_frame = tk.Frame(pending_canvas, bg="#1a1a2e")
        
        self.pending_phrases_frame.bind(
            "<Configure>",
            lambda e: pending_canvas.configure(scrollregion=pending_canvas.bbox("all"))
        )
        
        pending_canvas.create_window((0, 0), window=self.pending_phrases_frame, anchor="nw")
        pending_canvas.configure(yscrollcommand=pending_scrollbar.set)
        
        # Permettre le scroll avec la molette
        def _on_mousewheel(event):
            pending_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        pending_canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        pending_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        pending_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Label pour phrase vide
        self.empty_pending_label = tk.Label(
            self.pending_phrases_frame,
            text="Aucune phrase en attente.\nSaisissez une phrase ci-dessus et cliquez sur 'Ajouter à la liste'.",
            font=("Arial", 10, "italic"),
            bg="#1a1a2e",
            fg="#6b7280",
            justify=tk.CENTER
        )
        self.empty_pending_label.pack(pady=20)
        
        # Compteur de phrases en attente
        self.pending_count_label = tk.Label(
            main_frame,
            text="📊 Phrases en attente : 0",
            font=("Arial", 10, "bold"),
            bg="#1a1a2e",
            fg="#00ff88"
        )
        self.pending_count_label.pack(pady=(0, 10))
        
        # Frame pour les boutons d'action
        action_frame = tk.Frame(main_frame, bg="#1a1a2e")
        action_frame.pack(pady=(5, 0))
        
        # Bouton pour ajouter toutes les phrases au shop
        self.add_to_shop_btn = tk.Button(
            action_frame,
            text="🛒 AJOUTER CES PHRASES AU SHOP",
            command=lambda: self.add_all_phrases_to_shop(add_popup),
            bg="#8b5cf6",
            fg="white",
            font=("Arial", 12, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=12,
            activebackground="#7c3aed",
            borderwidth=0,
            state=tk.DISABLED
        )
        self.add_to_shop_btn.pack(side=tk.LEFT, padx=10)
        
        # Hover effect pour le bouton shop
        def shop_enter(e):
            if self.add_to_shop_btn['state'] != tk.DISABLED:
                self.add_to_shop_btn.config(bg="#7c3aed")
        def shop_leave(e):
            if self.add_to_shop_btn['state'] != tk.DISABLED:
                self.add_to_shop_btn.config(bg="#8b5cf6")
        self.add_to_shop_btn.bind("<Enter>", shop_enter)
        self.add_to_shop_btn.bind("<Leave>", shop_leave)
        
        # Bouton Fermer
        close_btn = tk.Button(
            action_frame,
            text="✖ Fermer",
            command=lambda: self.close_add_phrases_window(add_popup),
            bg="#6b7280",
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=12,
            activebackground="#4b5563",
            borderwidth=0
        )
        close_btn.pack(side=tk.LEFT, padx=10)
        
        # Hover effect
        close_btn.bind("<Enter>", lambda e: close_btn.config(bg="#4b5563"))
        close_btn.bind("<Leave>", lambda e: close_btn.config(bg="#6b7280"))
        
        # Info sur le nombre total de phrases actuel
        total_label = tk.Label(
            main_frame,
            text=f"📜 Total actuel dans le shop : {len(self.easter_messages)} phrases",
            font=("Arial", 9),
            bg="#1a1a2e",
            fg="#a78bfa"
        )
        total_label.pack(pady=(10, 0))
        self.total_shop_label = total_label
        
        # Stocker la référence du canvas pour unbind
        self.add_phrases_canvas = pending_canvas
        
        # Fermer proprement
        add_popup.protocol("WM_DELETE_WINDOW", lambda: self.close_add_phrases_window(add_popup))
    
    def add_phrase_to_temp_list(self, popup):
        """Ajouter une phrase à la liste temporaire"""
        phrase_text = self.new_phrase_entry.get("1.0", tk.END).strip()
        
        if not phrase_text:
            messagebox.showwarning(
                "Phrase vide",
                "Veuillez saisir une phrase avant de l'ajouter."
            )
            return
        
        if len(phrase_text) < 10:
            messagebox.showwarning(
                "Phrase trop courte",
                "La phrase doit contenir au moins 10 caractères."
            )
            return
        
        # Vérifier si la phrase existe déjà
        if phrase_text in self.easter_messages or phrase_text in self.new_phrases_list:
            messagebox.showwarning(
                "Phrase existante",
                "Cette phrase existe déjà dans le shop ou dans la liste en attente."
            )
            return
        
        # Ajouter à la liste temporaire
        self.new_phrases_list.append(phrase_text)
        
        # Cacher le label "aucune phrase"
        self.empty_pending_label.pack_forget()
        
        # Créer un widget pour afficher la nouvelle phrase
        phrase_idx = len(self.new_phrases_list) - 1
        phrase_frame = tk.Frame(self.pending_phrases_frame, bg="#16213e", padx=10, pady=8)
        phrase_frame.pack(fill=tk.X, pady=3, padx=5)
        
        # Numéro
        num_label = tk.Label(
            phrase_frame,
            text=f"#{phrase_idx + 1}",
            font=("Arial", 9, "bold"),
            bg="#16213e",
            fg="#ffd700",
            width=4
        )
        num_label.pack(side=tk.LEFT, padx=(0, 10))
        
        # Texte crypté (preview)
        crypted_preview = "🔒 " + "•" * min(len(phrase_text) // 2, 30) + "..."
        phrase_label = tk.Label(
            phrase_frame,
            text=crypted_preview,
            font=("Arial", 9),
            bg="#16213e",
            fg="#a78bfa",
            wraplength=400,
            justify=tk.LEFT,
            anchor="w"
        )
        phrase_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Bouton supprimer
        del_btn = tk.Button(
            phrase_frame,
            text="🗑️",
            command=lambda f=phrase_frame, p=phrase_text: self.remove_pending_phrase(f, p),
            bg="#dc2626",
            fg="white",
            font=("Arial", 9),
            cursor="hand2",
            relief=tk.FLAT,
            padx=8,
            pady=2,
            activebackground="#b91c1c",
            borderwidth=0
        )
        del_btn.pack(side=tk.RIGHT, padx=(10, 0))
        
        # Hover effect
        del_btn.bind("<Enter>", lambda e: del_btn.config(bg="#b91c1c"))
        del_btn.bind("<Leave>", lambda e: del_btn.config(bg="#dc2626"))
        
        # Vider la zone de saisie
        self.new_phrase_entry.delete("1.0", tk.END)
        
        # Mettre à jour le compteur
        self.update_pending_count()
        
        # Activer le bouton d'ajout au shop
        self.add_to_shop_btn.config(state=tk.NORMAL)
    
    def remove_pending_phrase(self, frame, phrase_text):
        """Supprimer une phrase de la liste en attente"""
        if phrase_text in self.new_phrases_list:
            self.new_phrases_list.remove(phrase_text)
        frame.destroy()
        
        # Mettre à jour le compteur
        self.update_pending_count()
        
        # Si plus de phrases, afficher le message vide et désactiver le bouton
        if not self.new_phrases_list:
            self.empty_pending_label.pack(pady=20)
            self.add_to_shop_btn.config(state=tk.DISABLED)
    
    def update_pending_count(self):
        """Mettre à jour le compteur de phrases en attente"""
        count = len(self.new_phrases_list)
        self.pending_count_label.config(text=f"📊 Phrases en attente : {count}")
    
    def add_all_phrases_to_shop(self, popup):
        """Ajouter toutes les phrases en attente au shop"""
        if not self.new_phrases_list:
            messagebox.showwarning(
                "Aucune phrase",
                "Aucune phrase en attente à ajouter."
            )
            return
        
        # Confirmation
        count = len(self.new_phrases_list)
        confirm = messagebox.askyesno(
            "Confirmer l'ajout",
            f"Voulez-vous ajouter {count} phrase{'s' if count > 1 else ''} au shop ?\n\n"
            "Ces phrases seront cryptées et disponibles à l'achat avec des crédits."
        )
        
        if not confirm:
            return
        
        # Ajouter les phrases à easter_messages
        added_count = 0
        for phrase in self.new_phrases_list:
            if phrase not in self.easter_messages:
                self.easter_messages.append(phrase)
                added_count += 1
        
        # Sauvegarder dans le fichier JSON (pour persistance)
        self.save_custom_phrases()
        
        # Message de succès
        messagebox.showinfo(
            "Phrases ajoutées !",
            f"✅ {added_count} phrase{'s' if added_count > 1 else ''} ajoutée{'s' if added_count > 1 else ''} au shop avec succès !\n\n"
            f"Total de phrases dans le shop : {len(self.easter_messages)}\n\n"
            "Les nouvelles phrases sont maintenant cryptées et disponibles à l'achat."
        )
        
        # Vider la liste temporaire et fermer
        self.new_phrases_list = []
        self.close_add_phrases_window(popup)
    
    def save_custom_phrases(self):
        """Sauvegarder les phrases personnalisées dans le fichier JSON"""
        try:
            # Charger les données existantes
            data = {}
            if os.path.exists(self.SAVE_FILE):
                with open(self.SAVE_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
            
            # Récupérer les crédits actuels
            credits = getattr(self, 'casino_state', {}).get('credits', self.saved_credits)
            
            # Sauvegarder avec les phrases personnalisées
            data['purchased_phrases'] = list(self.purchased_phrases)
            data['casino_credits'] = credits
            data['custom_phrases'] = self.easter_messages[46:]  # Les phrases ajoutées après les originales
            
            with open(self.SAVE_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            
            print(f"💾 Phrases personnalisées sauvegardées")
        except Exception as e:
            print(f"Erreur sauvegarde phrases personnalisées: {e}")
    
    def close_add_phrases_window(self, popup):
        """Fermer la fenêtre d'ajout de phrases"""
        try:
            if hasattr(self, 'add_phrases_canvas'):
                self.add_phrases_canvas.unbind_all("<MouseWheel>")
        except:
            pass
        popup.destroy()
    
    def show_missing_photos_alert(self, missing_count):
        """Afficher une alerte pour les photos manquantes (21ko)"""
        # Créer une fenêtre popup
        popup = tk.Toplevel(self.root)
        popup.title("⚠️ Photos manquantes détectées")
        popup.geometry("600x500")
        popup.configure(bg="white")
        popup.resizable(False, False)
        
        # Centrer la fenêtre
        popup.transient(self.root)
        popup.grab_set()
        
        # Frame principale
        main_frame = tk.Frame(popup, bg="white", padx=25, pady=25)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Titre
        title_label = tk.Label(
            main_frame,
            text="⚠️ Attention : Photos manquantes détectées",
            font=("Arial", 14, "bold"),
            bg="white",
            fg="#dc2626"
        )
        title_label.pack(pady=(0, 15))
        
        # Image portrait "inconnu"
        try:
            portrait_path = os.path.join(os.path.dirname(__file__), "assets", "portrait.png")
            if os.path.exists(portrait_path):
                portrait_img = Image.open(portrait_path)
                portrait_img = portrait_img.resize((100, 100), Image.Resampling.LANCZOS)
                self.portrait_photo = ImageTk.PhotoImage(portrait_img)
                
                portrait_label = tk.Label(
                    main_frame,
                    image=self.portrait_photo,
                    bg="white"
                )
                portrait_label.pack(pady=10)
        except Exception as e:
            print(f"Erreur chargement portrait: {e}")
        
        # Message principal
        message_text = (
            f"{missing_count} élève{'s' if missing_count > 1 else ''} sans photo {'ont' if missing_count > 1 else 'a'} été détecté{'s' if missing_count > 1 else ''}.\n\n"
            "Ces élèves ont une photo portrait \"inconnu\" pesant précisément 21 ko.\n\n"
            "Pour obtenir un meilleur trombinoscope, il est préférable de récupérer\n"
            "ces photos depuis une extraction Pronote si possible."
        )
        
        msg_label = tk.Label(
            main_frame,
            text=message_text,
            font=("Arial", 11),
            bg="white",
            wraplength=550,
            justify=tk.CENTER
        )
        msg_label.pack(pady=20)
        
        # Info supplémentaire
        info_label = tk.Label(
            main_frame,
            text="Ce message est purement informatif.",
            font=("Arial", 9, "italic"),
            bg="white",
            fg="#6b7280"
        )
        info_label.pack(pady=(10, 20))
        
        # Frame pour les boutons
        buttons_frame = tk.Frame(main_frame, bg="white")
        buttons_frame.pack(pady=5)
        
        # Bouton "Réparer avec une extraction Pronote" - Style moderne
        repair_container = tk.Frame(buttons_frame, bg="white")
        repair_container.pack(side=tk.LEFT, padx=5)
        
        repair_btn = tk.Button(
            repair_container,
            text="🔧 Réparer avec Pronote",
            command=lambda: [popup.destroy(), self.show_pronote_instructions()],
            bg=self.color_light_blue,
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=12,
            activebackground="#2563eb",
            borderwidth=0,
            highlightthickness=0
        )
        repair_btn.pack()
        
        # Effets hover pour repair_btn
        def repair_on_enter(e):
            repair_btn.config(bg="#2563eb")
        def repair_on_leave(e):
            repair_btn.config(bg=self.color_light_blue)
        repair_btn.bind("<Enter>", repair_on_enter)
        repair_btn.bind("<Leave>", repair_on_leave)
        
        # Bouton "Ignorer" - Style moderne
        ignore_container = tk.Frame(buttons_frame, bg="white")
        ignore_container.pack(side=tk.LEFT, padx=5)
        
        ignore_btn = tk.Button(
            ignore_container,
            text="Ignorer",
            command=popup.destroy,
            bg=self.color_green,
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=35,
            pady=12,
            activebackground="#047857",
            borderwidth=0,
            highlightthickness=0
        )
        ignore_btn.pack()
        
        # Effets hover pour ignore_btn
        def ignore_on_enter(e):
            ignore_btn.config(bg="#047857")
        def ignore_on_leave(e):
            ignore_btn.config(bg=self.color_green)
        ignore_btn.bind("<Enter>", ignore_on_enter)
        ignore_btn.bind("<Leave>", ignore_on_leave)
    
    def show_pronote_instructions(self):
        """Afficher les instructions pour exporter les photos depuis Pronote"""
        # Créer une fenêtre popup
        popup = tk.Toplevel(self.root)
        popup.title("📋 Instructions d'export Pronote")
        popup.geometry("700x650")
        popup.configure(bg="white")
        popup.resizable(False, False)
        
        # Centrer la fenêtre
        popup.transient(self.root)
        popup.grab_set()
        
        # Frame principale avec scrollbar
        main_frame = tk.Frame(popup, bg="white")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Titre
        title_label = tk.Label(
            main_frame,
            text="📋 Comment exporter les photos depuis Pronote",
            font=("Arial", 14, "bold"),
            bg="white",
            fg=self.color_blue
        )
        title_label.pack(pady=(0, 20))
        
        # Instructions
        instructions = [
            "1. Ouvrez Pronote.",
            "2. En haut, dans l'onglet \"Import-Export\", cliquez sur \"Exporter les photos des élèves\".",
            "3. Choisissez ou créez un dossier d'emplacement vide.",
            "4. Dans \"Syntaxe utilisée pour le nom des photos exportées\", choisir Nom-Prénom.",
            "   ⚠️ Le tiret est très important !"
        ]
        
        for instruction in instructions:
            instr_label = tk.Label(
                main_frame,
                text=instruction,
                font=("Arial", 10),
                bg="white",
                justify=tk.LEFT,
                anchor="w"
            )
            instr_label.pack(fill=tk.X, pady=3)
        
        # Afficher l'image d'exemple
        try:
            pronote_img_path = os.path.join(os.path.dirname(__file__), "assets", "pronote_export.png")
            if os.path.exists(pronote_img_path):
                pronote_img = Image.open(pronote_img_path)
                # Redimensionner l'image si nécessaire pour qu'elle tienne dans la fenêtre
                max_width = 650
                max_height = 300
                img_width, img_height = pronote_img.size
                
                # Calculer le ratio pour maintenir les proportions
                ratio = min(max_width / img_width, max_height / img_height)
                new_width = int(img_width * ratio)
                new_height = int(img_height * ratio)
                
                pronote_img = pronote_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                self.pronote_photo = ImageTk.PhotoImage(pronote_img)
                
                img_label = tk.Label(
                    main_frame,
                    image=self.pronote_photo,
                    bg="white",
                    relief=tk.SOLID,
                    borderwidth=1
                )
                img_label.pack(pady=15)
        except Exception as e:
            print(f"Erreur chargement image pronote: {e}")
        
        # Dernière instruction
        final_instr = tk.Label(
            main_frame,
            text="5. Lancez l'export.",
            font=("Arial", 10),
            bg="white",
            justify=tk.LEFT,
            anchor="w"
        )
        final_instr.pack(fill=tk.X, pady=3)
        
        # Séparateur
        separator = tk.Frame(main_frame, height=2, bg="#e5e7eb")
        separator.pack(fill=tk.X, pady=15)
        
        # Message d'information
        info_msg = tk.Label(
            main_frame,
            text="Une fois l'export terminé, cliquez sur \"Continuer\" pour sélectionner\nle dossier contenant les photos exportées.",
            font=("Arial", 9, "italic"),
            bg="white",
            fg="#6b7280",
            justify=tk.CENTER
        )
        info_msg.pack(pady=10)
        
        # Frame pour les boutons
        buttons_frame = tk.Frame(main_frame, bg="white")
        buttons_frame.pack(pady=10)
        
        # Bouton Annuler - Style moderne
        cancel_container = tk.Frame(buttons_frame, bg="white")
        cancel_container.pack(side=tk.LEFT, padx=5)
        
        cancel_btn = tk.Button(
            cancel_container,
            text="❌ Annuler",
            command=popup.destroy,
            bg="#6b7280",
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=12,
            activebackground="#4b5563",
            borderwidth=0,
            highlightthickness=0
        )
        cancel_btn.pack()
        
        def cancel_on_enter(e):
            cancel_btn.config(bg="#4b5563")
        def cancel_on_leave(e):
            cancel_btn.config(bg="#6b7280")
        cancel_btn.bind("<Enter>", cancel_on_enter)
        cancel_btn.bind("<Leave>", cancel_on_leave)
        
        # Bouton Continuer - Style moderne
        continue_container = tk.Frame(buttons_frame, bg="white")
        continue_container.pack(side=tk.LEFT, padx=5)
        
        continue_btn = tk.Button(
            continue_container,
            text="✅ Continuer",
            command=lambda: [popup.destroy(), self.repair_with_pronote()],
            bg=self.color_green,
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=12,
            activebackground="#047857",
            borderwidth=0,
            highlightthickness=0
        )
        continue_btn.pack()
        
        def continue_on_enter(e):
            continue_btn.config(bg="#047857")
        def continue_on_leave(e):
            continue_btn.config(bg=self.color_green)
        continue_btn.bind("<Enter>", continue_on_enter)
        continue_btn.bind("<Leave>", continue_on_leave)
    
    def repair_with_pronote(self):
        """Réparer les photos manquantes avec les photos Pronote"""
        # Demander le dossier Pronote
        pronote_folder = filedialog.askdirectory(
            title="Sélectionnez le dossier contenant les photos exportées de Pronote"
        )
        
        if not pronote_folder:
            return
        
        classe_jpg_path = self.classe_jpg_path.get()
        
        if not classe_jpg_path or not os.path.exists(classe_jpg_path):
            messagebox.showerror("Erreur", "Le dossier CLASSE_JPG n'est pas valide.")
            return
        
        self.update_status("Réparation des photos en cours...")
        
        # Statistiques
        photos_replaced = 0
        photos_still_missing = 0
        missing_photos_details = []
        
        try:
            # Créer un dictionnaire des photos Pronote (nom fichier -> chemin complet)
            pronote_photos = {}
            for file in os.listdir(pronote_folder):
                if file.lower().endswith(('.jpg', '.jpeg', '.png')):
                    # Extraire le nom sans extension
                    name = os.path.splitext(file)[0]
                    file_path = os.path.join(pronote_folder, file)
                    pronote_photos[name] = file_path
            
            # Parcourir les classes dans CLASSE_JPG
            for class_name in os.listdir(classe_jpg_path):
                class_path = os.path.join(classe_jpg_path, class_name)
                
                if not os.path.isdir(class_path):
                    continue
                
                # Parcourir les photos de la classe
                for file in os.listdir(class_path):
                    if file.lower().endswith(('.jpg', '.jpeg', '.png')):
                        file_path = os.path.join(class_path, file)
                        file_size = os.path.getsize(file_path)
                        
                        # Vérifier si c'est une photo "inconnu" de 21ko
                        if 20500 <= file_size <= 21500:
                            # Extraire le nom (NOM-Prenom)
                            name = os.path.splitext(file)[0]
                            
                            # Chercher la photo correspondante dans le dossier Pronote
                            if name in pronote_photos:
                                pronote_photo_path = pronote_photos[name]
                                
                                # Vérifier que la photo Pronote n'est pas aussi une photo "inconnu"
                                pronote_size = os.path.getsize(pronote_photo_path)
                                if 20500 <= pronote_size <= 21500:
                                    # La photo Pronote est aussi "inconnu"
                                    photos_still_missing += 1
                                    missing_photos_details.append({
                                        'name': name,
                                        'class': class_name,
                                        'reason': 'Photo Pronote aussi "inconnu"'
                                    })
                                else:
                                    # Remplacer la photo
                                    shutil.copy2(pronote_photo_path, file_path)
                                    photos_replaced += 1
                            else:
                                # Photo non trouvée dans le dossier Pronote
                                photos_still_missing += 1
                                missing_photos_details.append({
                                    'name': name,
                                    'class': class_name,
                                    'reason': 'Non trouvé dans dossier Pronote'
                                })
            
            # Afficher le récapitulatif
            self.show_repair_summary(photos_replaced, photos_still_missing, missing_photos_details)
            
            # Re-analyser les classes pour mettre à jour l'aperçu
            if photos_replaced > 0:
                self.analyze_classes()
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la réparation :\n{str(e)}")
            self.update_status("Erreur lors de la réparation")
            import traceback
            traceback.print_exc()
    
    def show_repair_summary(self, replaced_count, still_missing_count, missing_details):
        """Afficher le récapitulatif de la réparation"""
        # Créer une fenêtre popup
        popup = tk.Toplevel(self.root)
        popup.title("✅ Récapitulatif de la réparation")
        popup.geometry("650x550")
        popup.configure(bg="white")
        popup.resizable(False, False)
        
        # Centrer la fenêtre
        popup.transient(self.root)
        popup.grab_set()
        
        # Frame principale
        main_frame = tk.Frame(popup, bg="white", padx=25, pady=25)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Titre
        title_label = tk.Label(
            main_frame,
            text="✅ Récapitulatif de la réparation",
            font=("Arial", 14, "bold"),
            bg="white",
            fg=self.color_blue
        )
        title_label.pack(pady=(0, 20))
        
        # Statistiques
        stats_frame = tk.Frame(main_frame, bg="#f0f9ff", relief=tk.FLAT, padx=20, pady=15)
        stats_frame.pack(fill=tk.X, pady=10)
        
        # Photos remplacées
        replaced_label = tk.Label(
            stats_frame,
            text=f"✅ Photos correctement remplacées : {replaced_count}",
            font=("Arial", 11, "bold"),
            bg="#f0f9ff",
            fg=self.color_green
        )
        replaced_label.pack(anchor="w", pady=5)
        
        # Photos toujours manquantes
        missing_label = tk.Label(
            stats_frame,
            text=f"⚠️ Photos toujours manquantes : {still_missing_count}",
            font=("Arial", 11, "bold"),
            bg="#f0f9ff",
            fg="#dc2626" if still_missing_count > 0 else "#6b7280"
        )
        missing_label.pack(anchor="w", pady=5)
        
        # Détails des photos manquantes (si applicable)
        if still_missing_count > 0 and missing_details:
            details_label = tk.Label(
                main_frame,
                text="Détails des photos toujours manquantes :",
                font=("Arial", 10, "bold"),
                bg="white",
                fg=self.color_blue
            )
            details_label.pack(anchor="w", pady=(15, 5))
            
            # Zone de texte avec défilement pour les détails
            details_frame = tk.Frame(main_frame, bg="white")
            details_frame.pack(fill=tk.BOTH, expand=True, pady=5)
            
            details_text = scrolledtext.ScrolledText(
                details_frame,
                font=("Courier", 9),
                height=8,
                wrap=tk.WORD,
                bg="#f8fafc",
                relief=tk.FLAT
            )
            details_text.pack(fill=tk.BOTH, expand=True)
            
            for detail in missing_details[:20]:  # Limiter à 20 pour ne pas surcharger
                details_text.insert(tk.END, f"• {detail['name']} ({detail['class']}) - {detail['reason']}\n")
            
            if len(missing_details) > 20:
                details_text.insert(tk.END, f"\n... et {len(missing_details) - 20} autres\n")
            
            details_text.config(state=tk.DISABLED)
        
        # Message de conclusion
        if replaced_count > 0:
            conclusion_msg = "Les photos ont été remplacées avec succès dans CLASSE_JPG.\nVous pouvez maintenant générer le trombinoscope."
            conclusion_color = self.color_green
        else:
            conclusion_msg = "Aucune photo n'a pu être remplacée.\nVérifiez que le dossier Pronote contient des photos avec la bonne syntaxe (Nom-Prénom)."
            conclusion_color = "#dc2626"
        
        conclusion_label = tk.Label(
            main_frame,
            text=conclusion_msg,
            font=("Arial", 10),
            bg="white",
            fg=conclusion_color,
            justify=tk.CENTER
        )
        conclusion_label.pack(pady=15)
        
        # Frame pour les boutons
        buttons_frame = tk.Frame(main_frame, bg="white")
        buttons_frame.pack(pady=10)
        
        # Bouton Fermer - Style moderne
        close_container = tk.Frame(buttons_frame, bg="white")
        close_container.pack(side=tk.LEFT, padx=5)
        
        cancel_btn = tk.Button(
            close_container,
            text="✖ Fermer",
            command=popup.destroy,
            bg="#6b7280",
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=12,
            activebackground="#4b5563",
            borderwidth=0,
            highlightthickness=0
        )
        cancel_btn.pack()
        
        def close_on_enter(e):
            cancel_btn.config(bg="#4b5563")
        def close_on_leave(e):
            cancel_btn.config(bg="#6b7280")
        cancel_btn.bind("<Enter>", close_on_enter)
        cancel_btn.bind("<Leave>", close_on_leave)
        
        # Bouton Générer - Style moderne
        generate_container = tk.Frame(buttons_frame, bg="white")
        generate_container.pack(side=tk.LEFT, padx=5)
        
        generate_btn = tk.Button(
            generate_container,
            text="✨ Générer le trombinoscope",
            command=lambda: [popup.destroy(), self.generate_trombinoscope()],
            bg=self.color_green,
            fg="white",
            font=("Arial", 11, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=25,
            pady=12,
            activebackground="#047857",
            borderwidth=0,
            highlightthickness=0
        )
        generate_btn.pack()
        
        def generate_on_enter(e):
            generate_btn.config(bg="#047857")
        def generate_on_leave(e):
            generate_btn.config(bg=self.color_green)
        generate_btn.bind("<Enter>", generate_on_enter)
        generate_btn.bind("<Leave>", generate_on_leave)
        
        self.update_status(f"Réparation terminée : {replaced_count} photo(s) remplacée(s)")
        
    def setup_ui(self):
        """Configuration de l'interface utilisateur"""
        
        # Style
        style = ttk.Style()
        style.theme_use('clam')
        
        # Configuration du fond
        self.root.configure(bg=self.color_bg)
        
        # Définir l'icône de l'application avec Psyduck
        try:
            psyduck_icon_path = os.path.join(os.path.dirname(__file__), "assets", "psyduck.png")
            if os.path.exists(psyduck_icon_path):
                icon_img = Image.open(psyduck_icon_path)
                icon_photo = ImageTk.PhotoImage(icon_img)
                self.root.iconphoto(True, icon_photo)
        except Exception as e:
            print(f"Erreur lors du chargement de l'icône: {e}")
        
        # En-tête
        header_frame = tk.Frame(self.root, bg=self.color_blue, height=80)
        header_frame.pack(fill=tk.X, padx=0, pady=0)
        header_frame.pack_propagate(False)
        
        # Conteneur centré pour le titre et le logo
        title_container = tk.Frame(header_frame, bg=self.color_blue)
        title_container.place(relx=0.5, rely=0.5, anchor=tk.CENTER)
        
        # Titre
        title_label = tk.Label(
            title_container,
            text="📚 Générateur de Trombinoscope",
            font=("Arial", 24, "bold"),
            bg=self.color_blue,
            fg="white"
        )
        title_label.pack(side=tk.LEFT, padx=(0, 10))
        
        # Logo Psyduck décoratif dans le header - cliquable pour ouvrir le casino
        try:
            psyduck_path = os.path.join(os.path.dirname(__file__), "assets", "psyduck.png")
            psyduck_img = Image.open(psyduck_path)
            psyduck_img = psyduck_img.resize((50, 50), Image.Resampling.LANCZOS)
            self.psyduck_photo = ImageTk.PhotoImage(psyduck_img)
            
            psyduck_label = tk.Label(
                title_container,
                image=self.psyduck_photo,
                bg=self.color_blue,
                cursor="hand2"
            )
            psyduck_label.pack(side=tk.LEFT)
            # Clic pour ouvrir le jeu de casino
            psyduck_label.bind("<Button-1>", lambda e: self.show_casino_game())
        except Exception as e:
            print(f"Erreur chargement psyduck: {e}")
        
        # Frame principal
        main_frame = tk.Frame(self.root, bg=self.color_bg)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Section 1: Configuration
        config_frame = tk.LabelFrame(
            main_frame,
            text="📁 Configuration",
            font=("Arial", 13, "bold"),
            bg="white",
            fg=self.color_blue,
            padx=20,
            pady=20,
            relief=tk.FLAT,
            borderwidth=2,
            highlightthickness=1,
            highlightbackground="#e5e7eb"
        )
        config_frame.pack(fill=tk.X, pady=(0, 15))
        
        # Sélection du dossier
        tk.Label(
            config_frame,
            text="Dossier CLASSE_JPG :",
            font=("Arial", 10),
            bg="white"
        ).grid(row=0, column=0, sticky="w", pady=5)
        
        path_frame = tk.Frame(config_frame, bg="white")
        path_frame.grid(row=0, column=1, sticky="ew", padx=10)
        config_frame.columnconfigure(1, weight=1)
        
        tk.Entry(
            path_frame,
            textvariable=self.classe_jpg_path,
            font=("Arial", 10),
            width=50,
            relief=tk.FLAT,
            bg="#f8fafc",
            fg="#1e293b",
            insertbackground="#3b82f6"
        ).pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=6)
        
        # Bouton Parcourir moderne
        browse_container = tk.Frame(path_frame, bg="white")
        browse_container.pack(side=tk.LEFT, padx=(10, 0))
        
        browse_btn = tk.Button(
            browse_container,
            text="📁 Parcourir",
            command=self.select_folder,
            bg=self.color_green,
            fg="white",
            font=("Arial", 10, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=18,
            pady=8,
            activebackground="#047857",
            borderwidth=0,
            highlightthickness=0
        )
        browse_btn.pack()
        
        def browse_on_enter(e):
            browse_btn.config(bg="#047857")
        def browse_on_leave(e):
            browse_btn.config(bg=self.color_green)
        browse_btn.bind("<Enter>", browse_on_enter)
        browse_btn.bind("<Leave>", browse_on_leave)
        
        # Nom de l'établissement
        tk.Label(
            config_frame,
            text="Nom de l'établissement :",
            font=("Arial", 10),
            bg="white"
        ).grid(row=1, column=0, sticky="w", pady=5)
        
        tk.Entry(
            config_frame,
            textvariable=self.school_name,
            font=("Arial", 10),
            width=50,
            relief=tk.FLAT,
            bg="#f8fafc",
            fg="#1e293b",
            insertbackground="#3b82f6"
        ).grid(row=1, column=1, sticky="w", padx=10, ipady=6)
        
        # Année scolaire
        tk.Label(
            config_frame,
            text="Année scolaire :",
            font=("Arial", 10),
            bg="white"
        ).grid(row=2, column=0, sticky="w", pady=5)
        
        tk.Entry(
            config_frame,
            textvariable=self.school_year,
            font=("Arial", 10),
            width=50,
            relief=tk.FLAT,
            bg="#f8fafc",
            fg="#1e293b",
            insertbackground="#3b82f6"
        ).grid(row=2, column=1, sticky="w", padx=10, ipady=6)
        
        # Format de sortie
        tk.Label(
            config_frame,
            text="Format de sortie :",
            font=("Arial", 10),
            bg="white"
        ).grid(row=3, column=0, sticky="w", pady=5)
        
        format_frame = tk.Frame(config_frame, bg="white")
        format_frame.grid(row=3, column=1, sticky="w", padx=10)
        
        tk.Radiobutton(
            format_frame,
            text="Word (.docx)",
            variable=self.output_format,
            value="word",
            font=("Arial", 10),
            bg="white"
        ).pack(side=tk.LEFT, padx=(0, 20))
        
        tk.Radiobutton(
            format_frame,
            text="PDF (.pdf)",
            variable=self.output_format,
            value="pdf",
            font=("Arial", 10),
            bg="white"
        ).pack(side=tk.LEFT)
        
        # Section 2: Aperçu
        preview_frame = tk.LabelFrame(
            main_frame,
            text="📋 Aperçu des classes",
            font=("Arial", 13, "bold"),
            bg="white",
            fg=self.color_blue,
            padx=20,
            pady=20,
            relief=tk.FLAT,
            borderwidth=2,
            highlightthickness=1,
            highlightbackground="#e5e7eb"
        )
        preview_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        # Zone de texte avec défilement
        self.preview_text = scrolledtext.ScrolledText(
            preview_frame,
            font=("Courier", 9),
            height=10,
            wrap=tk.WORD,
            bg="#f8fafc",
            relief=tk.FLAT
        )
        self.preview_text.pack(fill=tk.BOTH, expand=True)
        
        # Section 3: Actions
        action_frame = tk.Frame(main_frame, bg=self.color_bg)
        action_frame.pack(fill=tk.X)
        
        # Bouton Analyser - Style moderne
        analyze_container = tk.Frame(action_frame, bg=self.color_bg)
        analyze_container.pack(side=tk.LEFT, padx=(0, 10))
        
        analyze_btn = tk.Button(
            analyze_container,
            text="🔍 Analyser les classes",
            command=self.analyze_classes,
            bg=self.color_light_blue,
            fg="white",
            font=("Arial", 12, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=30,
            pady=15,
            activebackground="#2563eb",
            borderwidth=0,
            highlightthickness=0
        )
        analyze_btn.pack()
        
        def analyze_on_enter(e):
            analyze_btn.config(bg="#2563eb")
        def analyze_on_leave(e):
            analyze_btn.config(bg=self.color_light_blue)
        analyze_btn.bind("<Enter>", analyze_on_enter)
        analyze_btn.bind("<Leave>", analyze_on_leave)
        
        # Bouton Générer - Style moderne
        generate_container = tk.Frame(action_frame, bg=self.color_bg)
        generate_container.pack(side=tk.LEFT)
        
        generate_btn = tk.Button(
            generate_container,
            text="✨ Générer le Trombinoscope",
            command=self.generate_trombinoscope,
            bg=self.color_green,
            fg="white",
            font=("Arial", 12, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=30,
            pady=15,
            activebackground="#047857",
            borderwidth=0,
            highlightthickness=0
        )
        generate_btn.pack()
        
        def generate_on_enter(e):
            generate_btn.config(bg="#047857")
        def generate_on_leave(e):
            generate_btn.config(bg=self.color_green)
        generate_btn.bind("<Enter>", generate_on_enter)
        generate_btn.bind("<Leave>", generate_on_leave)
        
        # GIF Psyduck à côté du bouton (cliquable pour le mini-jeu Easter Egg)
        psyduck_gif_container = tk.Frame(action_frame, bg=self.color_bg)
        psyduck_gif_container.pack(side=tk.LEFT, padx=(20, 0))
        
        self.psyduck_gif_label = tk.Label(psyduck_gif_container, bg=self.color_bg, cursor="hand2")
        self.psyduck_gif_label.pack()
        
        # Charger le GIF animé
        try:
            self.psyduck_gif_path = os.path.join(os.path.dirname(__file__), "assets", "psyduck_loading.gif")
            if os.path.exists(self.psyduck_gif_path):
                self.psyduck_gif = Image.open(self.psyduck_gif_path)
                self.psyduck_gif_frames = []
                try:
                    while True:
                        frame = self.psyduck_gif.copy()
                        frame = frame.resize((160, 80), Image.Resampling.LANCZOS)
                        self.psyduck_gif_frames.append(ImageTk.PhotoImage(frame))
                        self.psyduck_gif.seek(self.psyduck_gif.tell() + 1)
                except EOFError:
                    pass  # Fin du GIF
                self.psyduck_gif_index = 0
                # Afficher le premier frame du GIF
                if len(self.psyduck_gif_frames) > 0:
                    self.psyduck_gif_label.config(image=self.psyduck_gif_frames[0])
                    # Démarrer l'animation
                    self.animate_psyduck_gif()
                # Rendre le GIF cliquable pour ouvrir le mini-jeu
                self.psyduck_gif_label.bind("<Button-1>", lambda e: self.show_whack_a_psyduck_game())
        except Exception as e:
            print(f"Erreur chargement GIF Psyduck: {e}")
        
        # Barre de progression (initialement cachée)
        self.progress_frame = tk.Frame(main_frame, bg=self.color_bg)
        
        # Container horizontal pour le texte et la barre
        progress_content = tk.Frame(self.progress_frame, bg=self.color_bg)
        progress_content.pack(pady=(10, 5))
        
        self.progress_label = tk.Label(
            progress_content,
            text="",
            font=("Arial", 10, "bold"),
            bg=self.color_bg,
            fg=self.color_blue
        )
        self.progress_label.pack(pady=(0, 8))
        
        # Style moderne pour la barre de progression
        style.configure(
            "Modern.Horizontal.TProgressbar",
            troughcolor='#e5e7eb',
            background='#059669',
            darkcolor='#047857',
            lightcolor='#10b981',
            bordercolor='#d1d5db',
            thickness=20
        )
        
        self.progress_bar = ttk.Progressbar(
            progress_content,
            length=450,
            mode='determinate',
            style="Modern.Horizontal.TProgressbar"
        )
        self.progress_bar.pack()
        
        # Barre de statut
        self.status_label = tk.Label(
            self.root,
            text="Prêt à générer votre trombinoscope",
            font=("Arial", 9),
            bg=self.color_blue,
            fg="white",
            anchor="w",
            padx=10
        )
        self.status_label.pack(fill=tk.X, side=tk.BOTTOM)
    
    def animate_psyduck_gif(self):
        """Animer le GIF Psyduck en continu"""
        if hasattr(self, 'psyduck_gif_frames') and len(self.psyduck_gif_frames) > 0:
            # Vérifier si le label existe toujours et est visible
            if self.psyduck_gif_label.winfo_exists():
                self.psyduck_gif_index = (self.psyduck_gif_index + 1) % len(self.psyduck_gif_frames)
                self.psyduck_gif_label.config(image=self.psyduck_gif_frames[self.psyduck_gif_index])
                # Répéter l'animation toutes les 100ms
                self.root.after(100, self.animate_psyduck_gif)
    
    def select_folder(self):
        """Sélection du dossier CLASSE_JPG"""
        folder = filedialog.askdirectory(title="Sélectionner le dossier CLASSE_JPG")
        if folder:
            self.classe_jpg_path.set(folder)
            self.update_status(f"Dossier sélectionné : {folder}")
            
    def update_status(self, message):
        """Mise à jour de la barre de statut"""
        self.status_label.config(text=message)
        self.root.update_idletasks()
        
    def update_progress(self, current, total, message=""):
        """Mise à jour de la barre de progression avec message aléatoire"""
        if total > 0:
            progress_percent = (current / total) * 100
            self.progress_bar['value'] = progress_percent
            
            # Sélectionner un message aléatoire à chaque mise à jour
            random_msg = random.choice(self.loading_messages)
            
            # Afficher le message de progression avec le message aléatoire
            display_text = f"{message} ({current}/{total})\n💡 {random_msg}"
            self.progress_label.config(text=display_text)
            
            self.root.update_idletasks()
        
    def sort_class_name(self, class_name):
        """Tri intelligent des noms de classes"""
        # Extraction du type et du numéro
        
        # Secondes (2DE)
        if class_name.startswith('2DE'):
            number = int(re.search(r'\d+', class_name).group())
            return (1, number, class_name)
        
        # Premières générales (PG)
        elif class_name.startswith('PG'):
            number = int(re.search(r'\d+', class_name).group())
            return (2, number, class_name)
        
        # Premières STMG (PSTMG)
        elif class_name.startswith('PSTMG'):
            number = int(re.search(r'\d+', class_name).group())
            return (3, number, class_name)
        
        # Terminales générales (TG)
        elif class_name.startswith('TG') and not class_name.startswith('TGF'):
            number = int(re.search(r'\d+', class_name).group())
            return (4, number, class_name)
        
        # Terminales spéciales (TM, TGF, TRHC, etc.)
        elif class_name.startswith('TM') or class_name.startswith('TGF') or class_name.startswith('TRHC'):
            match = re.search(r'\d+', class_name)
            number = int(match.group()) if match else 0
            return (5, number, class_name)
        
        # BTS
        elif class_name.startswith('BTS'):
            # Extraction du type de BTS et du numéro
            bts_match = re.search(r'BTS\s*([A-Z]+)(\d+)', class_name)
            if bts_match:
                bts_type = bts_match.group(1)
                number = int(bts_match.group(2))
                return (6, ord(bts_type[0]), number, class_name)
            return (6, 0, 0, class_name)
        
        # Autres
        else:
            return (99, 0, class_name)
    
    def analyze_classes(self):
        """Analyse les classes dans le dossier sélectionné"""
        folder_path = self.classe_jpg_path.get()
        
        if not folder_path or not os.path.exists(folder_path):
            messagebox.showerror("Erreur", "Veuillez sélectionner un dossier valide.")
            return
        
        self.update_status("Analyse en cours...")
        self.preview_text.delete(1.0, tk.END)
        self.classes_data = {}
        
        try:
            # Parcours des sous-dossiers
            subdirs = [d for d in os.listdir(folder_path) 
                      if os.path.isdir(os.path.join(folder_path, d))]
            
            if not subdirs:
                messagebox.showwarning(
                    "Attention",
                    "Aucun sous-dossier de classe trouvé dans le dossier sélectionné."
                )
                return
            
            # Tri des classes
            subdirs.sort(key=self.sort_class_name)
            
            total_students = 0
            missing_photos_count = 0  # Compteur pour les photos manquantes (21ko)
            
            for class_name in subdirs:
                class_path = os.path.join(folder_path, class_name)
                
                # Recherche des fichiers JPG
                students = []
                for file in os.listdir(class_path):
                    if file.lower().endswith(('.jpg', '.jpeg', '.png')):
                        file_path = os.path.join(class_path, file)
                        # Vérifier la taille du fichier pour détecter les photos "inconnu" d'environ 21ko
                        file_size = os.path.getsize(file_path)
                        # Tolérance de ±500 bytes autour de 21ko pour détecter les photos "inconnu"
                        if 21000 <= file_size <= 22000:  # Entre 20.5ko et 21.5ko
                            missing_photos_count += 1
                        
                        # Extraction du nom (NOM-Prenom)
                        name = os.path.splitext(file)[0]
                        students.append({
                            'name': name,
                            'file_path': file_path
                        })
                
                # Tri alphabétique des élèves
                students.sort(key=lambda x: x['name'])
                
                if students:
                    self.classes_data[class_name] = students
                    total_students += len(students)
                    
                    # Affichage dans la prévisualisation
                    self.preview_text.insert(
                        tk.END,
                        f"📌 {class_name} ({len(students)} élèves)\n",
                        "class_header"
                    )
                    
                    for student in students[:5]:  # Afficher les 5 premiers
                        parts = student['name'].split('-')
                        if len(parts) == 2:
                            prenom = parts[1].capitalize()
                            nom = parts[0].upper()
                            self.preview_text.insert(tk.END, f"   • {prenom} {nom}\n")
                        else:
                            self.preview_text.insert(tk.END, f"   • {student['name']}\n")
                    
                    if len(students) > 5:
                        self.preview_text.insert(tk.END, f"   ... et {len(students) - 5} autres\n")
                    
                    self.preview_text.insert(tk.END, "\n")
            
            # Configuration des tags pour le texte
            self.preview_text.tag_config("class_header", foreground=self.color_blue, font=("Courier", 9, "bold"))
            
            # Résumé
            summary = f"\n{'='*60}\n"
            summary += f"✅ Total : {len(self.classes_data)} classes, {total_students} élèves\n"
            summary += f"{'='*60}\n"
            self.preview_text.insert(tk.END, summary, "summary")
            self.preview_text.tag_config("summary", foreground=self.color_green, font=("Courier", 9, "bold"))
            
            self.update_status(f"Analyse terminée : {len(self.classes_data)} classes trouvées")
            
            # Afficher l'alerte si des photos manquantes ont été détectées
            if missing_photos_count > 0:
                self.show_missing_photos_alert(missing_photos_count)
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'analyse :\n{str(e)}")
            self.update_status("Erreur lors de l'analyse")
    
    def generate_trombinoscope(self):
        """Génération du trombinoscope"""
        if not self.classes_data:
            messagebox.showwarning(
                "Attention",
                "Veuillez d'abord analyser les classes (bouton 'Analyser les classes')."
            )
            return
        
        # Cacher le GIF Psyduck lors du clic
        if hasattr(self, 'psyduck_gif_label'):
            self.psyduck_gif_label.pack_forget()
        
        # Demander où enregistrer le fichier
        output_format = self.output_format.get()
        file_extension = ".docx" if output_format == "word" else ".pdf"
        default_name = f"Trombinoscope_{self.school_year.get().replace('-', '_')}{file_extension}"
        
        output_file = filedialog.asksaveasfilename(
            defaultextension=file_extension,
            filetypes=[("Word Document", "*.docx"), ("PDF", "*.pdf")] if output_format == "word" else [("PDF", "*.pdf")],
            initialfile=default_name
        )
        
        if not output_file:
            # Réafficher le GIF si l'utilisateur annule
            if hasattr(self, 'psyduck_gif_label'):
                self.psyduck_gif_label.pack()
            return
        
        # Afficher la barre de progression
        self.progress_frame.pack(fill=tk.X, pady=(10, 0))
        self.progress_bar['value'] = 0
        
        self.update_status("Génération du trombinoscope en cours...")
        
        try:
            # Génération du document Word
            temp_docx = output_file if output_format == "word" else output_file.replace('.pdf', '_temp.docx')
            self.create_word_document(temp_docx)
            
            # Conversion en PDF si nécessaire
            if output_format == "pdf":
                self.update_progress(len(self.classes_data), len(self.classes_data) + 1, "Conversion en PDF")
                self.convert_to_pdf(temp_docx, output_file)
                os.remove(temp_docx)  # Suppression du fichier temporaire
            
            # Cacher la barre de progression
            self.progress_frame.pack_forget()
            
            # Réafficher le GIF après la génération
            if hasattr(self, 'psyduck_gif_label'):
                self.psyduck_gif_label.pack()
            
            self.update_status(f"✅ Trombinoscope généré avec succès : {output_file}")
            messagebox.showinfo(
                "Succès",
                f"Le trombinoscope a été généré avec succès !\n\nFichier : {output_file}"
            )
            
            # Ouvrir le dossier contenant le fichier
            if messagebox.askyesno("Ouvrir", "Voulez-vous ouvrir le dossier contenant le fichier ?"):
                os.startfile(os.path.dirname(output_file))
                
        except Exception as e:
            self.progress_frame.pack_forget()
            # Réafficher le GIF en cas d'erreur
            if hasattr(self, 'psyduck_gif_label'):
                self.psyduck_gif_label.pack()
            messagebox.showerror("Erreur", f"Erreur lors de la génération :\n{str(e)}")
            self.update_status("Erreur lors de la génération")
            import traceback
            traceback.print_exc()
    
    def create_word_document(self, output_file):
        """Création du document Word"""
        # Charger le document template comme base (pour conserver l'image de couverture)
        docx_path = os.path.join(os.path.dirname(__file__), "assets", "001 TROMBI COUV RECTO .docx")
        
        if os.path.exists(docx_path):
            # Utiliser le template comme base du document
            doc = Document(docx_path)
            self.update_progress(0, len(self.classes_data) + 1, "Chargement de la page de couverture")
            
            # Remplacer l'année scolaire dans la couverture
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # Remplacer différents formats d'année possibles
                    for old_year in ["2024-2025", "2024/2025", "2021-2022", "2020-2030"]:
                        if old_year in run.text:
                            run.text = run.text.replace(old_year, self.school_year.get())
        else:
            # Si le template n'existe pas, créer un document vide et ajouter une couverture par défaut
            doc = Document()
            self.add_cover_page(doc)
        
        # Configuration de la page en paysage avec marges de 1.5cm en haut et bas
        # Appliquer APRÈS le chargement du template pour ne pas écraser la première section
        for section in doc.sections:
            section.page_width = Inches(11.69)  # A4 paysage
            section.page_height = Inches(8.27)
            section.top_margin = Cm(1.5)  # Marges de 1.5cm en haut
            section.bottom_margin = Cm(1.5)  # Marges de 1.5cm en bas
            section.left_margin = Cm(0.7)
            section.right_margin = Cm(0.7)
        
        total_steps = len(self.classes_data) + 1
        current_step = 1
        
        # Pages des classes
        for idx, (class_name, students) in enumerate(self.classes_data.items()):
            self.update_progress(current_step, total_steps, f"Génération de la classe {class_name}")
            # Ne pas ajouter de page_break pour la première classe (idx == 0)
            # pour éviter une page blanche après la couverture
            if idx > 0:
                doc.add_page_break()
            self.add_class_page(doc, class_name, students)
            current_step += 1
        
        # Sauvegarde
        self.update_progress(total_steps, total_steps, "Sauvegarde du document")
        doc.save(output_file)
    
    def add_cover_page(self, doc):
        """Ajout de la page de couverture par défaut"""
        # En-tête
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = header.add_run(self.school_name.get())
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)  # Bleu foncé
        
        # Année scolaire
        year_para = doc.add_paragraph()
        year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = year_para.add_run(f"\nAnnée Scolaire {self.school_year.get()}")
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(5, 150, 105)  # Vert
        
        # Titre principal
        doc.add_paragraph("\n" * 3)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("TROMBINOSCOPE")
        run.font.size = Pt(48)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)
        
        # Informations supplémentaires
        doc.add_paragraph("\n" * 5)
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(f"{len(self.classes_data)} Classes")
        run.font.size = Pt(18)
        
        total_students = sum(len(students) for students in self.classes_data.values())
        info2 = doc.add_paragraph()
        info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info2.add_run(f"{total_students} Élèves")
        run.font.size = Pt(18)
    
    def add_class_page(self, doc, class_name, students):
        """Ajout d'une page de classe avec disposition 5 rangées × 9 colonnes (optimisé pour 40 élèves, max 45)"""
        
        # En-tête de la page avec le nom de la classe - ESPACEMENT MINIMAL
        class_header = doc.add_paragraph()
        class_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        class_header.paragraph_format.space_before = Pt(0)
        class_header.paragraph_format.space_after = Pt(2)  # Espacement minimal après le titre
        
        run = class_header.add_run(f"Classe {class_name}")
        run.font.size = Pt(16)  # Titre légèrement réduit pour gagner de l'espace
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)
        
        # Configuration de la grille: 5 rangées × 9 colonnes = 45 élèves max (optimisé pour 40)
        rows = 5
        cols = 9
        
        # Taille des photos optimisée pour 9 colonnes avec photos plus grandes
        photo_width = Cm(1.9)  # Augmenté à 1.9cm pour des photos plus grandes
        
        # Création du tableau avec toutes les lignes
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Configuration des marges du tableau pour maximiser l'espace
        for row in table.rows:
            for cell in row.cells:
                # Réduire les marges des cellules au MINIMUM
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    node = OxmlElement(f'w:{margin_name}')
                    node.set(qn('w:w'), '30')  # Marges augmentées à 30 pour 9 colonnes avec photos plus grandes
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                tcPr.append(tcMar)
        
        # Remplissage du tableau - FORCER à remplir TOUTES les lignes
        for row_idx in range(rows):
            for col_idx in range(cols):
                idx = row_idx * cols + col_idx
                
                if idx < len(students):
                    student = students[idx]
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Ajout de la photo
                    try:
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Réduire l'espacement du paragraphe
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        
                        run = paragraph.add_run()
                        run.add_picture(student['file_path'], width=photo_width)
                        
                        # Ajout du nom
                        name_parts = student['name'].split('-')
                        if len(name_parts) == 2:
                            prenom = name_parts[1].capitalize()
                            nom = name_parts[0].upper()
                            display_name = f"{prenom}\n{nom}"
                        else:
                            display_name = student['name']
                        
                        name_para = cell.add_paragraph()
                        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        name_para.paragraph_format.space_before = Pt(0)
                        name_para.paragraph_format.space_after = Pt(0)
                        
                        run = name_para.add_run(display_name)
                        run.font.size = Pt(7)  # Augmenté à 7pt pour meilleure lisibilité avec photos plus grandes
                        run.font.bold = True
                        
                    except Exception as e:
                        # En cas d'erreur, afficher juste le nom
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(f"[Photo manquante]\n{student['name']}")
                        run.font.size = Pt(7)
        
        # Si la classe a plus de 45 élèves, afficher un avertissement
        if len(students) > rows * cols:
            doc.add_paragraph()
            warning = doc.add_paragraph()
            warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = warning.add_run(f"⚠ {len(students) - (rows * cols)} élèves supplémentaires non affichés (limite: {rows * cols} élèves/page)")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(220, 38, 38)  # Rouge
    
    def convert_to_pdf(self, docx_file, pdf_file):
        """Conversion du document Word en PDF"""
        try:
            from docx2pdf import convert
            convert(docx_file, pdf_file)
        except ImportError:
            messagebox.showwarning(
                "Attention",
                "La conversion en PDF nécessite l'installation de 'docx2pdf'.\n"
                "Le fichier Word a été généré à la place."
            )
            # Renommer le fichier .docx en gardant le nom demandé
            new_name = pdf_file.replace('.pdf', '.docx')
            shutil.move(docx_file, new_name)
        except Exception as e:
            messagebox.showerror(
                "Erreur de conversion",
                f"Impossible de convertir en PDF.\n{str(e)}\n\n"
                "Le fichier Word a été conservé."
            )


def main():
    """Point d'entrée de l'application"""
    root = tk.Tk()
    app = TrombinoscopeApp(root)
    
    # Sauvegarder les crédits du casino à la fermeture de l'application
    def on_closing():
        # Sauvegarder les données du casino avant de fermer
        if hasattr(app, 'casino_state') and app.casino_state:
            app.save_purchased_phrases()
        root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.mainloop()


if __name__ == "__main__":
    main()
